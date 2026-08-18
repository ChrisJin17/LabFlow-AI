import streamlit as st
from datetime import date, datetime, timedelta
import sqlite3
import json
import base64
from supabase import create_client, Client
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Image
)


def add_inventory_item(
    chemical_name,
    category,
    unit,
    current_stock,
    minimum_stock,
    cost_per_unit,
    supplier,
    location
):

    inventory_data = {
        "chemical_name": chemical_name,
        "category": category,
        "unit": unit,
        "current_stock": current_stock,
        "minimum_stock": minimum_stock,
        "cost_per_unit": cost_per_unit,
        "supplier": supplier,
        "location": location,
        "last_updated": str(date.today())
    }

    (
        supabase
        .table("inventory")
        .insert(inventory_data)
        .execute()
    )

def get_inventory():

    response = (
        supabase
        .table("inventory")
        .select(
            "id, "
            "chemical_name, "
            "category, "
            "unit, "
            "current_stock, "
            "minimum_stock, "
            "cost_per_unit, "
            "supplier, "
            "location, "
            "last_updated"
        )
        .order(
            "chemical_name"
        )
        .execute()
    )

    inventory = []

    for item in response.data:

        inventory.append((
            item["id"],
            item["chemical_name"],
            item["category"],
            item["unit"],
            item["current_stock"],
            item["minimum_stock"],
            item["cost_per_unit"],
            item["supplier"],
            item["location"],
            item["last_updated"]
        ))

    return inventory

# --------------------------------------------------
# INVENTORY CHECK AND DEDUCTION
# --------------------------------------------------

def get_inventory_item_by_name(chemical_name):

    response = (
        supabase
        .table("inventory")
        .select(
            "id, chemical_name, category, unit, "
            "current_stock, minimum_stock, cost_per_unit"
        )
        .ilike(
            "chemical_name",
            chemical_name.strip()
        )
        .limit(1)
        .execute()
    )

    if not response.data:
        return None

    item = response.data[0]

    return (
        item["id"],
        item["chemical_name"],
        item["category"],
        item["unit"],
        item["current_stock"],
        item["minimum_stock"],
        item["cost_per_unit"]
    )


def check_inventory_amount(chemical_name, required_amount):

    if not chemical_name or required_amount <= 0:
        return True, None

    item = get_inventory_item_by_name(chemical_name)

    # Chemical is not managed in inventory.
    # Do not block experiment creation.
    if item is None:
        return True, None

    current_stock = item[4]

    if current_stock < required_amount:
        return False, {
            "chemical": item[1],
            "required": required_amount,
            "available": current_stock,
            "unit": item[3]
        }

    return True, None


def deduct_inventory(
    chemical_name,
    amount
):

    if not chemical_name or amount <= 0:
        return

    item = get_inventory_item_by_name(
        chemical_name
    )

    if item is None:
        return

    inventory_id = item[0]
    current_stock = item[4] or 0

    new_stock = current_stock - amount

    (
        supabase
        .table("inventory")
        .update({
            "current_stock": new_stock,
            "last_updated": str(date.today())
        })
        .eq(
            "id",
            inventory_id
        )
        .execute()
    )

def restore_inventory(
    chemical_name,
    amount
):

    if not chemical_name or amount <= 0:
        return

    item = get_inventory_item_by_name(
        chemical_name
    )

    if item is None:
        return

    inventory_id = item[0]
    current_stock = item[4] or 0

    new_stock = current_stock + amount

    (
        supabase
        .table("inventory")
        .update({
            "current_stock": new_stock,
            "last_updated": str(date.today())
        })
        .eq(
            "id",
            inventory_id
        )
        .execute()
    )

def restore_experiment_inventory(
    reagents_json,
    solvent,
    solvent_volume
):

    try:
        reagents = json.loads(
            reagents_json
        )
    except:
        reagents = []

    # Restore reagents
    for reagent in reagents:

        chemical_name = reagent.get(
            "Chemical",
            ""
        )

        required_mass_mg = reagent.get(
            "Required mass (mg)",
            0
        ) or 0

        if not chemical_name:
            continue

        required_mass_g = (
            required_mass_mg / 1000
        )

        restore_inventory(
            chemical_name,
            required_mass_g
        )

    # Restore solvent
    if solvent and solvent_volume > 0:

        restore_inventory(
            solvent,
            solvent_volume
        )

def check_experiment_inventory(
    calculated_reagents,
    solvent,
    solvent_volume
):

    shortages = []

    # Check reagents
    for reagent in calculated_reagents:

        chemical_name = reagent.get("Chemical", "")
        required_mass_mg = reagent.get(
            "Required mass (mg)",
            0
        )

        if not chemical_name:
            continue

        # Inventory stores reagent mass in grams
        required_mass_g = required_mass_mg / 1000

        enough, shortage = check_inventory_amount(
            chemical_name,
            required_mass_g
        )

        if not enough:
            shortages.append(shortage)

    # Check solvent
    if solvent and solvent_volume > 0:

        enough, shortage = check_inventory_amount(
            solvent,
            solvent_volume
        )

        if not enough:
            shortages.append(shortage)

    return shortages

def deduct_experiment_inventory(
    calculated_reagents,
    solvent,
    solvent_volume
):

    # Deduct reagents
    for reagent in calculated_reagents:

        chemical_name = reagent.get("Chemical", "")
        required_mass_mg = reagent.get(
            "Required mass (mg)",
            0
        )

        if not chemical_name:
            continue

        required_mass_g = required_mass_mg / 1000

        deduct_inventory(
            chemical_name,
            required_mass_g
        )

    # Deduct solvent
    if solvent and solvent_volume > 0:

        deduct_inventory(
            solvent,
            solvent_volume
        )

def mark_inventory_deducted(
    experiment_id
):

    (
        supabase
        .table("experiments")
        .update({
            "inventory_deducted": 1
        })
        .eq(
            "id",
            experiment_id
        )
        .execute()
    )

def feature_card(icon, title, description):

    card_html = f"""
<div style="border:1px solid rgba(128,128,128,0.20); border-radius:16px; padding:20px; min-height:155px; margin-bottom:14px; background:rgba(128,128,128,0.035);">
<div style="font-size:1.55rem; margin-bottom:8px;">{icon}</div>
<div style="font-size:1.1rem; font-weight:700; margin-bottom:7px;">{title}</div>
<div style="font-size:0.92rem; line-height:1.55; opacity:0.72;">{description}</div>
</div>
"""

    st.markdown(
        card_html,
        unsafe_allow_html=True
    )

def demo_data_exists():

    if (
        "user" not in st.session_state
        or st.session_state.user is None
    ):
        return False

    response = (
        supabase
        .table("experiments")
        .select("id")
        .eq(
            "researcher",
            "Demo User"
        )
        .eq(
            "user_id",
            str(st.session_state.user.id)
        )
        .limit(1)
        .execute()
    )

    return bool(response.data)

def load_demo_data():

  

    demo_experiments = [
        {
            "name": "Catalyst Screen A",
            "researcher": "Demo User",
            "date": str(date.today() - timedelta(days=4)),
            "starting_material": "4-Bromoacetophenone",
            "sm_mw": 199.04,
            "sm_mass": 199.04,
            "sm_mmol": 1.000,
            "reagents": [
                {
                    "Chemical": "Phenylboronic acid",
                    "Role": "Reagent",
                    "MW (g/mol)": 121.93,
                    "Equiv": 1.5,
                    "Required mmol": 1.500,
                    "Required mass (mg)": 182.90,
                    "Estimated Cost (HKD)": 91.45
                },
                {
                    "Chemical": "Pd(PPh3)4",
                    "Role": "Catalyst",
                    "MW (g/mol)": 1155.56,
                    "Equiv": 0.02,
                    "Required mmol": 0.020,
                    "Required mass (mg)": 23.11,
                    "Estimated Cost (HKD)": 69.33
                }
            ],
            "solvent": "THF",
            "solvent_volume": 10.0,
            "temperature": 80.0,
            "reaction_time": 12.0,
            "yield": 35.0,
            "observation": (
                "The reaction mixture remained pale yellow. "
                "TLC after 2 h showed substantial starting material."
            ),
            "objective": "Evaluate low catalyst loading.",
            "status": "Failed",
            "key_result": "Low catalyst loading gave poor conversion and 35% yield.",
            "next_step": "Increase catalyst loading to 5 mol%.",
            "solvent_cost_per_l": 100.0,
            "estimated_total_cost": 161.78,
            "ai_procedure": None
        },
        {
            "name": "Catalyst Screen B",
            "researcher": "Demo User",
            "date": str(date.today() - timedelta(days=3)),
            "starting_material": "4-Bromoacetophenone",
            "sm_mw": 199.04,
            "sm_mass": 199.04,
            "sm_mmol": 1.000,
            "reagents": [
                {
                    "Chemical": "Phenylboronic acid",
                    "Role": "Reagent",
                    "MW (g/mol)": 121.93,
                    "Equiv": 1.5,
                    "Required mmol": 1.500,
                    "Required mass (mg)": 182.90,
                    "Estimated Cost (HKD)": 91.45
                },
                {
                    "Chemical": "Pd(PPh3)4",
                    "Role": "Catalyst",
                    "MW (g/mol)": 1155.56,
                    "Equiv": 0.05,
                    "Required mmol": 0.050,
                    "Required mass (mg)": 57.78,
                    "Estimated Cost (HKD)": 173.34
                }
            ],
            "solvent": "THF",
            "solvent_volume": 10.0,
            "temperature": 80.0,
            "reaction_time": 12.0,
            "yield": 68.0,
            "observation": (
                "The mixture turned dark brown after 20 min. "
                "TLC showed improved but incomplete conversion."
            ),
            "objective": "Evaluate whether 5 mol% catalyst improves conversion.",
            "status": "Partial Success",
            "key_result": "Yield improved from 35% to 68%.",
            "next_step": "Increase temperature to 90 °C.",
            "solvent_cost_per_l": 100.0,
            "estimated_total_cost": 265.79,
            "ai_procedure": None
        },
        {
            "name": "Temperature Optimization",
            "researcher": "Demo User",
            "date": str(date.today() - timedelta(days=2)),
            "starting_material": "4-Bromoacetophenone",
            "sm_mw": 199.04,
            "sm_mass": 199.04,
            "sm_mmol": 1.000,
            "reagents": [
                {
                    "Chemical": "Phenylboronic acid",
                    "Role": "Reagent",
                    "MW (g/mol)": 121.93,
                    "Equiv": 1.5,
                    "Required mmol": 1.500,
                    "Required mass (mg)": 182.90,
                    "Estimated Cost (HKD)": 91.45
                },
                {
                    "Chemical": "Pd(PPh3)4",
                    "Role": "Catalyst",
                    "MW (g/mol)": 1155.56,
                    "Equiv": 0.05,
                    "Required mmol": 0.050,
                    "Required mass (mg)": 57.78,
                    "Estimated Cost (HKD)": 173.34
                }
            ],
            "solvent": "THF",
            "solvent_volume": 10.0,
            "temperature": 90.0,
            "reaction_time": 12.0,
            "yield": 82.0,
            "observation": (
                "TLC showed almost complete consumption of starting material."
            ),
            "objective": "Evaluate the effect of higher temperature.",
            "status": "Successful",
            "key_result": "Increasing temperature to 90 °C improved yield to 82%.",
            "next_step": "Test solvent effects under the optimized conditions.",
            "solvent_cost_per_l": 100.0,
            "estimated_total_cost": 265.79,
            "ai_procedure": None
        },
        {
            "name": "Solvent Screen DCM",
            "researcher": "Demo User",
            "date": str(date.today() - timedelta(days=1)),
            "starting_material": "4-Bromoacetophenone",
            "sm_mw": 199.04,
            "sm_mass": 199.04,
            "sm_mmol": 1.000,
            "reagents": [
                {
                    "Chemical": "Phenylboronic acid",
                    "Role": "Reagent",
                    "MW (g/mol)": 121.93,
                    "Equiv": 1.5,
                    "Required mmol": 1.500,
                    "Required mass (mg)": 182.90,
                    "Estimated Cost (HKD)": 91.45
                }
            ],
            "solvent": "DCM",
            "solvent_volume": 10.0,
            "temperature": 40.0,
            "reaction_time": 12.0,
            "yield": 48.0,
            "observation": "Incomplete conversion was observed by TLC.",
            "objective": "Compare DCM with THF under related conditions.",
            "status": "Partial Success",
            "key_result": "DCM gave lower yield than the optimized THF condition.",
            "next_step": "Continue using THF for further optimization.",
            "solvent_cost_per_l": 120.0,
            "estimated_total_cost": 92.65,
            "ai_procedure": None
        },
        {
            "name": "Optimized Reaction",
            "researcher": "Demo User",
            "date": str(date.today()),
            "starting_material": "4-Bromoacetophenone",
            "sm_mw": 199.04,
            "sm_mass": 199.04,
            "sm_mmol": 1.000,
            "reagents": [
                {
                    "Chemical": "Phenylboronic acid",
                    "Role": "Reagent",
                    "MW (g/mol)": 121.93,
                    "Equiv": 1.5,
                    "Required mmol": 1.500,
                    "Required mass (mg)": 182.90,
                    "Estimated Cost (HKD)": 91.45
                },
                {
                    "Chemical": "Pd(PPh3)4",
                    "Role": "Catalyst",
                    "MW (g/mol)": 1155.56,
                    "Equiv": 0.05,
                    "Required mmol": 0.050,
                    "Required mass (mg)": 57.78,
                    "Estimated Cost (HKD)": 173.34
                }
            ],
            "solvent": "THF",
            "solvent_volume": 10.0,
            "temperature": 90.0,
            "reaction_time": 16.0,
            "yield": 88.0,
            "observation": (
                "TLC showed complete consumption of starting material. "
                "The desired product was isolated in high yield."
            ),
            "objective": "Validate the optimized reaction conditions.",
            "status": "Successful",
            "key_result": "The optimized conditions gave 88% isolated yield.",
            "next_step": "Evaluate substrate scope.",
            "solvent_cost_per_l": 100.0,
            "estimated_total_cost": 265.79,
            "ai_procedure": None
        }
    ]

    for exp in demo_experiments:
        st.write(
            "DEBUG USER ID:",
            st.session_state.user.id
        )

        experiment_data = {
            "user_id": str(st.session_state.user.id),
            "experiment_name": exp["name"],
            "researcher": exp["researcher"],
            "experiment_date": exp["date"],
            "starting_material": exp["starting_material"],
            "starting_material_mw": exp["sm_mw"],
            "starting_material_mass": exp["sm_mass"],
            "starting_material_mmol": exp["sm_mmol"],
            "reagents": exp["reagents"],
            "solvent": exp["solvent"],
            "solvent_volume": exp["solvent_volume"],
            "temperature": exp["temperature"],
            "reaction_time": exp["reaction_time"],
            "yield_percent": exp["yield"],
            "observation": exp["observation"],
            "objective": exp["objective"],
            "status": exp["status"],
            "key_result": exp["key_result"],
            "next_step": exp["next_step"],
            "solvent_cost_per_l": exp["solvent_cost_per_l"],
            "estimated_total_cost": exp["estimated_total_cost"],
            "ai_procedure": exp["ai_procedure"],
            "inventory_deducted": 0
        }

        (
            supabase
            .table("experiments")
            .insert(
                experiment_data
            )
            .execute()
        )



def clear_demo_data():

    (
        supabase
        .table("experiments")
        .delete()
        .eq(
            "researcher",
            "Demo User"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .execute()
    )

    

def save_experiment(
    experiment_name,
    researcher,
    experiment_date,
    sm_name,
    sm_mw,
    sm_mass,
    sm_mmol,
    reagents,
    solvent,
    solvent_volume,
    temperature,
    reaction_time,
    yield_percent,
    observation,
    objective,
    status,
    key_result,
    next_step,
    solvent_cost_per_l,
    estimated_total_cost
):
    if (
        "user" not in st.session_state
        or st.session_state.user is None
    ):
        raise Exception(
            "No authenticated user found."
        )
    
    experiment_data = {
        "user_id": st.session_state.user.id,
        "experiment_name": experiment_name,
        "researcher": researcher,
        "experiment_date": str(experiment_date),

        "starting_material": sm_name,
        "starting_material_mw": sm_mw,
        "starting_material_mass": sm_mass,
        "starting_material_mmol": sm_mmol,

        "reagents": reagents,

        "solvent": solvent,
        "solvent_volume": solvent_volume,

        "temperature": temperature,
        "reaction_time": reaction_time,

        "yield_percent": yield_percent,
        "observation": observation,

        "objective": objective,
        "status": status,
        "key_result": key_result,
        "next_step": next_step,

        "solvent_cost_per_l": solvent_cost_per_l,
        "estimated_total_cost": estimated_total_cost,

        "inventory_deducted": 0
    }

    response = (
        supabase
        .table("experiments")
        .insert(experiment_data)
        .execute()
    )

    if not response.data:
        raise Exception(
            "Supabase did not return the saved experiment."
        )

    experiment_id = response.data[0]["id"]

    return experiment_id

def get_all_experiments():

    response = (
        supabase
        .table("experiments")
        .select(
            "id, "
            "experiment_name, "
            "researcher, "
            "experiment_date, "
            "starting_material, "
            "solvent, "
            "solvent_volume, "
            "yield_percent"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .order(
            "experiment_date",
            desc=True
        )
        .order(
            "id",
            desc=True
        )
        .execute()
    )

    experiments = []

    for exp in response.data:

        experiments.append((
            exp["id"],
            exp["experiment_name"],
            exp["researcher"],
            exp["experiment_date"],
            exp["starting_material"],
            exp["solvent"],
            exp["solvent_volume"],
            exp["yield_percent"]
        ))

    return experiments

def delete_experiment(experiment_id):

    # --------------------------------------------------
    # GET EXPERIMENT DATA BEFORE DELETION
    # --------------------------------------------------

    response = (
        supabase
        .table("experiments")
        .select(
            "id, reagents, solvent, "
            "solvent_volume, inventory_deducted"
        )
        .eq(
            "id",
            experiment_id
        )
        .limit(1)
        .execute()
    )

    if not response.data:
        return False

    experiment = response.data[0]

    reagents = experiment.get(
        "reagents",
        []
    ) or []

    solvent = experiment.get(
        "solvent",
        ""
    )

    solvent_volume = experiment.get(
        "solvent_volume",
        0
    ) or 0

    inventory_deducted = experiment.get(
        "inventory_deducted",
        0
    ) or 0

    # --------------------------------------------------
    # RESTORE INVENTORY
    # --------------------------------------------------

    if inventory_deducted == 1:

        for reagent in reagents:

            chemical_name = reagent.get(
                "Chemical",
                ""
            )

            required_mass_mg = reagent.get(
                "Required mass (mg)",
                0
            ) or 0

            if chemical_name:

                required_mass_g = (
                    required_mass_mg / 1000
                )

                restore_inventory(
                    chemical_name,
                    required_mass_g
                )

        if solvent and solvent_volume > 0:

            restore_inventory(
                solvent,
                solvent_volume
            )

    # --------------------------------------------------
    # DELETE EXPERIMENT
    # --------------------------------------------------

    (
        supabase
        .table("experiments")
        .delete()
        .eq(
            "id",
            experiment_id
        )
        .execute()
    )

    return True

def get_experiments_for_report():

    response = (
        supabase
        .table("experiments")
        .select(
            "id, "
            "experiment_name, "
            "researcher, "
            "experiment_date, "
            "starting_material, "
            "starting_material_mass, "
            "starting_material_mmol, "
            "reagents, "
            "solvent, "
            "solvent_volume, "
            "temperature, "
            "reaction_time, "
            "yield_percent, "
            "observation, "
            "objective, "
            "status, "
            "key_result, "
            "next_step"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .order(
            "experiment_date"
        )
        .order(
            "id"
        )
        .execute()
    )

    experiments = []

    for exp in response.data:

        experiments.append((
            exp["id"],
            exp["experiment_name"],
            exp["researcher"],
            exp["experiment_date"],
            exp["starting_material"],
            exp["starting_material_mass"],
            exp["starting_material_mmol"],

            # Keep old report code compatible
            json.dumps(
                exp["reagents"] or []
            ),

            exp["solvent"],
            exp["solvent_volume"],
            exp["temperature"],
            exp["reaction_time"],
            exp["yield_percent"],
            exp["observation"],
            exp["objective"],
            exp["status"],
            exp["key_result"],
            exp["next_step"]
        ))

    return experiments

def get_experiments_for_ai_search():

    response = (
        supabase
        .table("experiments")
        .select(
            "id, "
            "experiment_name, "
            "researcher, "
            "experiment_date, "
            "starting_material, "
            "starting_material_mass, "
            "starting_material_mmol, "
            "reagents, "
            "solvent, "
            "solvent_volume, "
            "temperature, "
            "reaction_time, "
            "yield_percent, "
            "observation, "
            "objective, "
            "status, "
            "key_result, "
            "next_step, "
            "ai_procedure"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )       
        .order(
            "experiment_date"
        )
        .order(
            "id"
        )
        .execute()
    )

    experiments = []

    for exp in response.data:

        experiments.append((
            exp["id"],
            exp["experiment_name"],
            exp["researcher"],
            exp["experiment_date"],
            exp["starting_material"],
            exp["starting_material_mass"],
            exp["starting_material_mmol"],

            json.dumps(
                exp["reagents"] or []
            ),

            exp["solvent"],
            exp["solvent_volume"],
            exp["temperature"],
            exp["reaction_time"],
            exp["yield_percent"],
            exp["observation"],
            exp["objective"],
            exp["status"],
            exp["key_result"],
            exp["next_step"],
            exp["ai_procedure"]
        ))

    return experiments

def answer_research_question(question, experiments):

    experiment_text = ""

    for exp in experiments:

        experiment_text += f"""
Experiment ID: EXP-{exp[0]:04d}
Experiment Name: {exp[1]}
Researcher: {exp[2]}
Date: {exp[3]}

Starting Material: {exp[4]}
Starting Material Mass: {exp[5]} mg
Starting Material Amount: {exp[6]} mmol

Reagents: {exp[7]}

Solvent: {exp[8]}
Solvent Volume: {exp[9]} mL

Temperature: {exp[10]} °C
Reaction Time: {exp[11]} h

Yield: {exp[12]} %

Observation:
{exp[13]}

Objective:
{exp[14]}

Status:
{exp[15]}

Key Result:
{exp[16]}

Next Step:
{exp[17]}

Experimental Procedure:
{exp[18]}

-----------------------------------
"""

    prompt = f"""
You are LabFlow AI, a chemistry research memory assistant.

Answer the user's question based ONLY on the experimental
records provided below.

User question:
{question}

Important rules:

- Do not invent experiments, values, chemicals or conclusions.
- Use only information contained in the experimental records.
- When relevant, mention experiment IDs such as EXP-0012.
- When comparing experiments, clearly state the recorded
  conditions and results.
- If the database does not contain enough information,
  say so clearly.
- Distinguish recorded facts from your interpretation.
- Do not claim causation unless it is directly supported by
  the experimental records.
- Use concise professional scientific English.
- Where useful, identify trends across multiple experiments.
- If asked for a suggested next step, base it only on recorded
  results and explicitly label it as a suggestion.

Experimental records:

{experiment_text}
"""

    return cloud_ai_chat(
        prompt
    )

import re

def extract_experiment_ids(text):

    matches = re.findall(
        r"EXP-(\d+)",
        text
    )

    experiment_ids = []

    for match in matches:

        experiment_id = int(match)

        if experiment_id not in experiment_ids:
            experiment_ids.append(
                experiment_id
            )

    return experiment_ids

def filter_experiments_for_search(
    experiments,
    start_date=None,
    end_date=None,
    researcher_filter="All",
    status_filter="All",
    solvent_filter="All"
):

    filtered = []

    for exp in experiments:

        try:
            exp_date = datetime.strptime(
                exp[3],
                "%Y-%m-%d"
            ).date()
        except:
            continue

        # Date filter
        if start_date and exp_date < start_date:
            continue

        if end_date and exp_date > end_date:
            continue

        # Researcher filter
        if (
            researcher_filter != "All"
            and exp[2] != researcher_filter
        ):
            continue

        # Status filter
        if (
            status_filter != "All"
            and (exp[15] or "Not specified")
            != status_filter
        ):
            continue

        # Solvent filter
        if (
            solvent_filter != "All"
            and (exp[8] or "Not recorded")
            != solvent_filter
        ):
            continue

        filtered.append(exp)

    return filtered

def get_experiment_by_id(experiment_id):

    response = (
        supabase
        .table("experiments")
        .select(
            "id, "
            "experiment_name, "
            "researcher, "
            "experiment_date, "
            "starting_material, "
            "starting_material_mw, "
            "starting_material_mass, "
            "starting_material_mmol, "
            "reagents, "
            "solvent, "
            "solvent_volume, "
            "temperature, "
            "reaction_time, "
            "yield_percent, "
            "observation, "
            "objective, "
            "status, "
            "key_result, "
            "next_step, "
            "estimated_total_cost, "
            "ai_procedure"
        )
        .eq(
            "id",
            experiment_id
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .limit(1)
        .execute()
    )

    if not response.data:
        return None

    exp = response.data[0]

    # Keep the original SQLite-style tuple structure
    # so existing LabFlow code does not need to change.
    experiment = (
        exp["id"],
        exp["experiment_name"],
        exp["researcher"],
        exp["experiment_date"],
        exp["starting_material"],
        exp["starting_material_mw"],
        exp["starting_material_mass"],
        exp["starting_material_mmol"],

        # Supabase stores reagents as JSONB.
        # Convert back to JSON string for old LabFlow code.
        json.dumps(
            exp["reagents"] or []
        ),

        exp["solvent"],
        exp["solvent_volume"],
        exp["temperature"],
        exp["reaction_time"],
        exp["yield_percent"],
        exp["observation"],
        exp["objective"],
        exp["status"],
        exp["key_result"],
        exp["next_step"],
        exp["estimated_total_cost"],
        exp["ai_procedure"]
    )

    return experiment

def save_ai_procedure(
    experiment_id,
    ai_procedure
):

    (
        supabase
        .table("experiments")
        .update({
            "ai_procedure": ai_procedure
        })
        .eq(
            "id",
            experiment_id
        )
        .execute()
    )



def save_attachment(
    experiment_id,
    uploaded_file,
    attachment_category,
    notes
):

    file_bytes = uploaded_file.getvalue()

    file_data_base64 = base64.b64encode(
        file_bytes
    ).decode("utf-8")

    attachment_data = {
        "experiment_id": experiment_id,
        "file_name": uploaded_file.name,
        "file_type": uploaded_file.type,
        "attachment_category": attachment_category,
        "file_data": file_data_base64,
        "notes": notes,
        "uploaded_at": datetime.now().isoformat()
    }

    (
        supabase
        .table("attachments")
        .insert(
            attachment_data
        )
        .execute()
    )
    
def get_attachments(
    experiment_id
):

    response = (
        supabase
        .table("attachments")
        .select(
            "id, "
            "file_name, "
            "file_type, "
            "attachment_category, "
            "file_data, "
            "notes, "
            "uploaded_at"
        )
        .eq(
            "experiment_id",
            experiment_id
        )
        .order(
            "id",
            desc=True
        )
        .execute()
    )

    attachments = []

    for item in response.data:

        try:

            file_bytes = base64.b64decode(
                item["file_data"]
            )

        except Exception:

            file_bytes = b""

        attachments.append((
            item["id"],
            item["file_name"],
            item["file_type"],
            item["attachment_category"],
            file_bytes,
            item["notes"],
            item["uploaded_at"]
        ))

    return attachments

def generate_word_lab_note(experiment, ai_procedure=None):

    (
        experiment_id,
        experiment_name,
        researcher,
        experiment_date,
        starting_material,
        sm_mw,
        sm_mass,
        sm_mmol,
        reagents_json,
        solvent,
        solvent_volume,
        temperature,
        reaction_time,
        yield_percent,
        observation,
        objective,
        status,
        key_result,
        next_step,
        estimated_total_cost,
        saved_ai_procedure
    ) = experiment

    attachments = get_attachments(
        experiment_id
    )

    try:
        reagents = json.loads(reagents_json)
    except:
        reagents = []

    document = Document()

    # --------------------------------------------------
    # TITLE
    # --------------------------------------------------

    title = document.add_heading(
        "LabFlow AI Laboratory Record",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle.add_run(
        f"EXP-{experiment_id:04d} — {experiment_name}"
    ).bold = True

    document.add_paragraph("")

    # --------------------------------------------------
    # BASIC INFORMATION
    # --------------------------------------------------

    document.add_heading(
        "Experiment Information",
        level=1
    )

    table = document.add_table(
        rows=4,
        cols=2
    )

    table.style = "Table Grid"

    table.cell(0, 0).text = "Experiment ID"
    table.cell(0, 1).text = f"EXP-{experiment_id:04d}"

    table.cell(1, 0).text = "Researcher"
    table.cell(1, 1).text = researcher or ""

    table.cell(2, 0).text = "Date"
    table.cell(2, 1).text = experiment_date or ""

    table.cell(3, 0).text = "Status"
    table.cell(3, 1).text = status or "Not specified"

    # --------------------------------------------------
    # OBJECTIVE
    # --------------------------------------------------

    document.add_heading(
        "Objective",
        level=1
    )

    document.add_paragraph(
        objective or "No objective recorded."
    )

    # --------------------------------------------------
    # REACTION SETUP
    # --------------------------------------------------

    document.add_heading(
        "Reaction Setup",
        level=1
    )

    document.add_heading(
        "Starting Material",
        level=2
    )

    sm_table = document.add_table(
        rows=2,
        cols=5
    )

    sm_table.style = "Table Grid"

    headers = [
        "Chemical",
        "MW (g/mol)",
        "Mass (mg)",
        "Amount (mmol)",
        "Equiv"
    ]

    for i, header in enumerate(headers):
        sm_table.cell(0, i).text = header

    sm_table.cell(1, 0).text = starting_material or ""
    sm_table.cell(1, 1).text = f"{sm_mw:.2f}"
    sm_table.cell(1, 2).text = f"{sm_mass:.2f}"
    sm_table.cell(1, 3).text = f"{sm_mmol:.3f}"
    sm_table.cell(1, 4).text = "1.00"

    # --------------------------------------------------
    # REAGENTS
    # --------------------------------------------------

    document.add_heading(
        "Reagents",
        level=2
    )

    valid_reagents = [
        r for r in reagents
        if r.get("Chemical")
    ]

    if valid_reagents:

        reagent_table = document.add_table(
            rows=1,
            cols=6
        )

        reagent_table.style = "Table Grid"

        reagent_headers = [
            "Chemical",
            "Role",
            "MW",
            "Mass (mg)",
            "mmol",
            "Equiv"
        ]

        for i, header in enumerate(reagent_headers):
            reagent_table.cell(0, i).text = header

        for reagent in valid_reagents:

            row = reagent_table.add_row().cells

            row[0].text = str(
                reagent.get("Chemical", "")
            )

            row[1].text = str(
                reagent.get("Role", "Reagent")
            )

            row[2].text = (
                f"{reagent.get('MW (g/mol)', 0):.2f}"
            )

            row[3].text = (
                f"{reagent.get('Required mass (mg)', 0):.2f}"
            )

            row[4].text = (
                f"{reagent.get('Required mmol', 0):.3f}"
            )

            row[5].text = (
                f"{reagent.get('Equiv', 0):.3f}"
            )

    else:

        document.add_paragraph(
            "No reagents recorded."
        )

    # --------------------------------------------------
    # REACTION CONDITIONS
    # --------------------------------------------------

    document.add_heading(
        "Reaction Conditions",
        level=1
    )

    conditions_table = document.add_table(
        rows=3,
        cols=2
    )

    conditions_table.style = "Table Grid"

    conditions_table.cell(0, 0).text = "Solvent"
    conditions_table.cell(
        0,
        1
    ).text = (
        f"{solvent or 'Not recorded'} "
        f"({solvent_volume or 0:.2f} mL)"
    )

    conditions_table.cell(
        1,
        0
    ).text = "Temperature"

    conditions_table.cell(
        1,
        1
    ).text = (
        f"{temperature or 0:.1f} °C"
    )

    conditions_table.cell(
        2,
        0
    ).text = "Reaction Time"

    conditions_table.cell(
        2,
        1
    ).text = (
        f"{reaction_time or 0:.1f} h"
    )

    # --------------------------------------------------
    # OBSERVATIONS
    # --------------------------------------------------

    document.add_heading(
        "Experimental Observations",
        level=1
    )

    document.add_paragraph(
        observation or "No observations recorded."
    )

    # --------------------------------------------------
    # AI PROFESSIONAL PROCEDURE
    # --------------------------------------------------

    document.add_heading(
        "Experimental Procedure",
        level=1
    )

    if ai_procedure:

        document.add_paragraph(
            ai_procedure
        )

    else:

        document.add_paragraph(
            "No AI-generated experimental procedure available."
        )


    # --------------------------------------------------
    # RESULT
    # --------------------------------------------------

    document.add_heading(
        "Result",
        level=1
    )

    result_table = document.add_table(
        rows=3,
        cols=2
    )

    result_table.style = "Table Grid"

    result_table.cell(
        0,
        0
    ).text = "Isolated Yield"

    if yield_percent is not None:
        result_table.cell(
            0,
            1
        ).text = f"{yield_percent:.1f}%"
    else:
        result_table.cell(
            0,
            1
        ).text = "Not recorded"

    result_table.cell(
        1,
        0
    ).text = "Status"

    result_table.cell(
        1,
        1
    ).text = status or "Not specified"

    result_table.cell(
        2,
        0
    ).text = "Estimated Cost"

    if estimated_total_cost is not None:
        result_table.cell(
            2,
            1
        ).text = f"HK${estimated_total_cost:.2f}"
    else:
        result_table.cell(
            2,
            1
        ).text = "Not recorded"

    # --------------------------------------------------
    # KEY RESULT
    # --------------------------------------------------

    document.add_heading(
        "Key Result",
        level=1
    )

    document.add_paragraph(
        key_result or "No key result recorded."
    )

    # --------------------------------------------------
    # ATTACHMENTS
    # --------------------------------------------------

    document.add_heading(
        "Experimental Attachments",
        level=1
    )

    if attachments:

        attachment_table = document.add_table(
            rows=1,
            cols=4
        )

        attachment_table.style = "Table Grid"

        attachment_headers = [
            "Category",
            "File Name",
            "Notes",
            "Uploaded"
        ]

        for i, header in enumerate(
            attachment_headers
        ):
            attachment_table.cell(
                0,
                i
            ).text = header

        for attachment in attachments:

            row = attachment_table.add_row().cells

            row[0].text = str(
                attachment[3] or ""
            )

            row[1].text = str(
                attachment[1] or ""
            )

            row[2].text = str(
                attachment[5] or ""
            )

            row[3].text = str(
                attachment[6] or ""
            )

    else:

        document.add_paragraph(
            "No experimental attachments recorded."
        )

    # --------------------------------------------------
    # FONT
    # --------------------------------------------------

    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)

    # --------------------------------------------------
    # SAVE TO MEMORY
    # --------------------------------------------------

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer

def generate_pdf_lab_note(experiment, ai_procedure=None):

    (
        experiment_id,
        experiment_name,
        researcher,
        experiment_date,
        starting_material,
        sm_mw,
        sm_mass,
        sm_mmol,
        reagents_json,
        solvent,
        solvent_volume,
        temperature,
        reaction_time,
        yield_percent,
        observation,
        objective,
        status,
        key_result,
        next_step,
        estimated_total_cost,
        saved_ai_procedure
    ) = experiment

    attachments = get_attachments(
        experiment_id
    )

    try:
        reagents = json.loads(reagents_json)
    except:
        reagents = []

    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "LabFlowTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=8
    )

    subtitle_style = ParagraphStyle(
        "LabFlowSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=14
    )

    heading_style = ParagraphStyle(
        "LabFlowHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        "LabFlowBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        spaceAfter=6
    )

    story = []

    # --------------------------------------------------
    # TITLE
    # --------------------------------------------------

    story.append(
        Paragraph(
            "LabFlow AI Laboratory Record",
            title_style
        )
    )

    story.append(
        Paragraph(
            f"EXP-{experiment_id:04d} - {experiment_name}",
            subtitle_style
        )
    )

    # --------------------------------------------------
    # EXPERIMENT INFORMATION
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Experiment Information",
            heading_style
        )
    )

    info_data = [
        ["Experiment ID", f"EXP-{experiment_id:04d}"],
        ["Researcher", researcher or ""],
        ["Date", experiment_date or ""],
        ["Status", status or "Not specified"]
    ]

    info_table = Table(
        info_data,
        colWidths=[45 * mm, 110 * mm]
    )

    info_table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5)
        ])
    )

    story.append(info_table)
    story.append(Spacer(1, 8))

    # --------------------------------------------------
    # OBJECTIVE
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Objective",
            heading_style
        )
    )

    story.append(
        Paragraph(
            objective or "No objective recorded.",
            body_style
        )
    )

    # --------------------------------------------------
    # STARTING MATERIAL
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Reaction Setup",
            heading_style
        )
    )

    story.append(
        Paragraph(
            "Starting Material",
            styles["Heading3"]
        )
    )

    sm_data = [
        [
            "Chemical",
            "MW (g/mol)",
            "Mass (mg)",
            "mmol",
            "Equiv"
        ],
        [
            starting_material or "",
            f"{sm_mw:.2f}",
            f"{sm_mass:.2f}",
            f"{sm_mmol:.3f}",
            "1.00"
        ]
    ]

    sm_table = Table(
        sm_data,
        colWidths=[
            50 * mm,
            25 * mm,
            25 * mm,
            25 * mm,
            20 * mm
        ]
    )

    sm_table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("VALIGN", (0, 0), (-1, -1), "TOP")
        ])
    )

    story.append(sm_table)
    story.append(Spacer(1, 8))

    # --------------------------------------------------
    # REAGENTS
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Reagents",
            styles["Heading3"]
        )
    )

    valid_reagents = [
        r for r in reagents
        if r.get("Chemical")
    ]

    if valid_reagents:

        reagent_data = [
            [
                "Chemical",
                "Role",
                "MW",
                "Mass (mg)",
                "mmol",
                "Equiv"
            ]
        ]

        for reagent in valid_reagents:

            reagent_data.append([
                str(reagent.get("Chemical", "")),
                str(reagent.get("Role", "Reagent")),
                f"{reagent.get('MW (g/mol)', 0):.2f}",
                f"{reagent.get('Required mass (mg)', 0):.2f}",
                f"{reagent.get('Required mmol', 0):.3f}",
                f"{reagent.get('Equiv', 0):.3f}"
            ])

        reagent_table = Table(
            reagent_data,
            colWidths=[
                42 * mm,
                25 * mm,
                22 * mm,
                25 * mm,
                22 * mm,
                20 * mm
            ],
            repeatRows=1
        )

        reagent_table.setStyle(
            TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP")
            ])
        )

        story.append(reagent_table)

    else:

        story.append(
            Paragraph(
                "No reagents recorded.",
                body_style
            )
        )

    # --------------------------------------------------
    # CONDITIONS
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Reaction Conditions",
            heading_style
        )
    )

    conditions_data = [
        [
            "Solvent",
            (
                f"{solvent or 'Not recorded'} "
                f"({solvent_volume or 0:.2f} mL)"
            )
        ],
        [
            "Temperature",
            f"{temperature or 0:.1f} °C"
        ],
        [
            "Reaction Time",
            f"{reaction_time or 0:.1f} h"
        ]
    ]

    conditions_table = Table(
        conditions_data,
        colWidths=[45 * mm, 110 * mm]
    )

    conditions_table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9)
        ])
    )

    story.append(conditions_table)

    # --------------------------------------------------
    # OBSERVATIONS
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Experimental Observations",
            heading_style
        )
    )

    story.append(
        Paragraph(
            observation or "No observations recorded.",
            body_style
        )
    )

    # --------------------------------------------------
    # AI PROFESSIONAL PROCEDURE
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Experimental Procedure",
            heading_style
        )
    )

    story.append(
        Paragraph(
            (
                ai_procedure
                if ai_procedure
                else "No AI-generated experimental procedure available."
            ),
            body_style
        )
    )

    # --------------------------------------------------
    # RESULT
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Result",
            heading_style
        )
    )

    result_yield = (
        f"{yield_percent:.1f}%"
        if yield_percent is not None
        else "Not recorded"
    )

    result_cost = (
        f"HK${estimated_total_cost:.2f}"
        if estimated_total_cost is not None
        else "Not recorded"
    )

    result_data = [
        ["Isolated Yield", result_yield],
        ["Status", status or "Not specified"],
        ["Estimated Cost", result_cost]
    ]

    result_table = Table(
        result_data,
        colWidths=[45 * mm, 110 * mm]
    )

    result_table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9)
        ])
    )

    story.append(result_table)

    # --------------------------------------------------
    # ATTACHMENTS
    # --------------------------------------------------

    story.append(
        Paragraph(
            "Experimental Attachments",
            heading_style
        )
    )

    if attachments:

        attachment_data = [
            [
                "Category",
                "File Name",
                "Notes",
                "Uploaded"
            ]
        ]

        for attachment in attachments:

            attachment_data.append([
                str(attachment[3] or ""),
                str(attachment[1] or ""),
                str(attachment[5] or ""),
                str(attachment[6] or "")
            ])

        attachment_table = Table(
            attachment_data,
            colWidths=[
                28 * mm,
                45 * mm,
                55 * mm,
                30 * mm
            ],
            repeatRows=1
        )

        attachment_table.setStyle(
            TableStyle([
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    colors.grey
                ),
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.lightgrey
                ),
                (
                    "FONTNAME",
                    (0, 0),
                    (-1, 0),
                    "Helvetica-Bold"
                ),
                (
                    "FONTSIZE",
                    (0, 0),
                    (-1, -1),
                    7.5
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "TOP"
                )
            ])
        )

        story.append(
            attachment_table
        )

        # --------------------------------------------------
        # IMAGE ATTACHMENTS
        # --------------------------------------------------

        image_attachments = [
            attachment
            for attachment in attachments
            if attachment[2] in [
                "image/png",
                "image/jpeg"
            ]
        ]

        if image_attachments:

            story.append(
                Spacer(1, 10)
            )

            story.append(
                Paragraph(
                    "Attachment Images",
                    heading_style
                )
            )

            for attachment in image_attachments:

                file_name = attachment[1]
                category = attachment[3]
                file_data = attachment[4]
                notes = attachment[5]

                image_buffer = BytesIO(
                    file_data
                )

                story.append(
                    Paragraph(
                        f"<b>{category}</b> — {file_name}",
                        body_style
                    )
                )

                if notes:

                    story.append(
                        Paragraph(
                            notes,
                            body_style
                        )
                    )

                try:

                    pdf_image = Image(
                        image_buffer
                    )

                    max_width = 150 * mm
                    max_height = 100 * mm

                    image_width = pdf_image.imageWidth
                    image_height = pdf_image.imageHeight

                    scale = min(
                        max_width / image_width,
                        max_height / image_height,
                        1
                    )

                    pdf_image.drawWidth = (
                        image_width * scale
                    )

                    pdf_image.drawHeight = (
                        image_height * scale
                    )

                    story.append(
                        pdf_image
                    )

                    story.append(
                        Spacer(1, 10)
                    )

                except Exception:

                    story.append(
                        Paragraph(
                            "Image preview could not be generated.",
                            body_style
                        )
                    )
    else:

        story.append(
            Paragraph(
                "No experimental attachments recorded.",
                body_style
            )
        )

    doc.build(story)

    buffer.seek(0)

    return buffer

    # --------------------------------------------------
    # TITLE
    # --------------------------------------------------

    title = document.add_heading(
        "LabFlow AI Laboratory Record",
        level=0
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()

    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle.add_run(
        f"EXP-{experiment_id:04d} — {experiment_name}"
    ).bold = True

    document.add_paragraph("")

    # --------------------------------------------------
    # BASIC INFORMATION
    # --------------------------------------------------

    document.add_heading(
        "Experiment Information",
        level=1
    )

    table = document.add_table(
        rows=4,
        cols=2
    )

    table.style = "Table Grid"

    table.cell(0, 0).text = "Experiment ID"
    table.cell(0, 1).text = f"EXP-{experiment_id:04d}"

    table.cell(1, 0).text = "Researcher"
    table.cell(1, 1).text = researcher or ""

    table.cell(2, 0).text = "Date"
    table.cell(2, 1).text = experiment_date or ""

    table.cell(3, 0).text = "Status"
    table.cell(3, 1).text = status or "Not specified"

    # --------------------------------------------------
    # OBJECTIVE
    # --------------------------------------------------

    document.add_heading(
        "Objective",
        level=1
    )

    document.add_paragraph(
        objective or "No objective recorded."
    )

    # --------------------------------------------------
    # REACTION SETUP
    # --------------------------------------------------

    document.add_heading(
        "Reaction Setup",
        level=1
    )

    document.add_heading(
        "Starting Material",
        level=2
    )

    sm_table = document.add_table(
        rows=2,
        cols=5
    )

    sm_table.style = "Table Grid"

    headers = [
        "Chemical",
        "MW (g/mol)",
        "Mass (mg)",
        "Amount (mmol)",
        "Equiv"
    ]

    for i, header in enumerate(headers):
        sm_table.cell(0, i).text = header

    sm_table.cell(1, 0).text = starting_material or ""
    sm_table.cell(1, 1).text = f"{sm_mw:.2f}"
    sm_table.cell(1, 2).text = f"{sm_mass:.2f}"
    sm_table.cell(1, 3).text = f"{sm_mmol:.3f}"
    sm_table.cell(1, 4).text = "1.00"

    # --------------------------------------------------
    # REAGENTS
    # --------------------------------------------------

    document.add_heading(
        "Reagents",
        level=2
    )

    valid_reagents = [
        r for r in reagents
        if r.get("Chemical")
    ]

    if valid_reagents:

        reagent_table = document.add_table(
            rows=1,
            cols=6
        )

        reagent_table.style = "Table Grid"

        reagent_headers = [
            "Chemical",
            "Role",
            "MW",
            "Mass (mg)",
            "mmol",
            "Equiv"
        ]

        for i, header in enumerate(
            reagent_headers
        ):
            reagent_table.cell(
                0,
                i
            ).text = header

        for reagent in valid_reagents:

            row = reagent_table.add_row().cells

            row[0].text = str(
                reagent.get(
                    "Chemical",
                    ""
                )
            )

            row[1].text = str(
                reagent.get(
                    "Role",
                    "Reagent"
                )
            )

            row[2].text = (
                f"{reagent.get('MW (g/mol)', 0):.2f}"
            )

            row[3].text = (
                f"{reagent.get('Required mass (mg)', 0):.2f}"
            )

            row[4].text = (
                f"{reagent.get('Required mmol', 0):.3f}"
            )

            row[5].text = (
                f"{reagent.get('Equiv', 0):.3f}"
            )

    else:

        document.add_paragraph(
            "No reagents recorded."
        )

    # --------------------------------------------------
    # REACTION CONDITIONS
    # --------------------------------------------------

    document.add_heading(
        "Reaction Conditions",
        level=1
    )

    conditions_table = document.add_table(
        rows=3,
        cols=2
    )

    conditions_table.style = "Table Grid"

    conditions_table.cell(
        0,
        0
    ).text = "Solvent"

    conditions_table.cell(
        0,
        1
    ).text = (
        f"{solvent or 'Not recorded'} "
        f"({solvent_volume or 0:.2f} mL)"
    )

    conditions_table.cell(
        1,
        0
    ).text = "Temperature"

    conditions_table.cell(
        1,
        1
    ).text = (
        f"{temperature or 0:.1f} °C"
    )

    conditions_table.cell(
        2,
        0
    ).text = "Reaction Time"

    conditions_table.cell(
        2,
        1
    ).text = (
        f"{reaction_time or 0:.1f} h"
    )

    # --------------------------------------------------
    # OBSERVATIONS
    # --------------------------------------------------

    document.add_heading(
        "Experimental Observations",
        level=1
    )

    document.add_paragraph(
        observation
        or "No observations recorded."
    )

    # --------------------------------------------------
    # RESULT
    # --------------------------------------------------

    document.add_heading(
        "Result",
        level=1
    )

    result_table = document.add_table(
        rows=3,
        cols=2
    )

    result_table.style = "Table Grid"

    result_table.cell(
        0,
        0
    ).text = "Isolated Yield"

    if yield_percent is not None:
        result_table.cell(
            0,
            1
        ).text = f"{yield_percent:.1f}%"

    else:
        result_table.cell(
            0,
            1
        ).text = "Not recorded"

    result_table.cell(
        1,
        0
    ).text = "Status"

    result_table.cell(
        1,
        1
    ).text = status or "Not specified"

    result_table.cell(
        2,
        0
    ).text = "Estimated Cost"

    if estimated_total_cost is not None:
        result_table.cell(
            2,
            1
        ).text = (
            f"HK${estimated_total_cost:.2f}"
        )
    else:
        result_table.cell(
            2,
            1
        ).text = "Not recorded"

    # --------------------------------------------------
    # KEY RESULT
    # --------------------------------------------------

    document.add_heading(
        "Key Result",
        level=1
    )

    document.add_paragraph(
        key_result
        or "No key result recorded."
    )

    # --------------------------------------------------
    # NEXT STEP
    # --------------------------------------------------

    document.add_heading(
        "Next Step",
        level=1
    )

    document.add_paragraph(
        next_step
        or "No next step recorded."
    )

    # --------------------------------------------------
    # FONT
    # --------------------------------------------------

    styles = document.styles

    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    # --------------------------------------------------
    # SAVE TO MEMORY
    # --------------------------------------------------

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer


def generate_ai_report(experiments, report_type):

    experiment_text = ""

    for exp in experiments:

        experiment_text += f"""
Experiment ID: EXP-{exp[0]:04d}
Experiment Name: {exp[1]}
Researcher: {exp[2]}
Date: {exp[3]}

Starting Material: {exp[4]}
Starting Material Mass: {exp[5]} mg
Starting Material Amount: {exp[6]} mmol

Reagents: {exp[7]}

Solvent: {exp[8]}
Solvent Volume: {exp[9]} mL
Temperature: {exp[10]} °C
Reaction Time: {exp[11]} h

Yield: {exp[12]} %

Observation:
{exp[13]}

Objective:
{exp[14]}

Status:
{exp[15]}

Key Result:
{exp[16]}

Next Step:
{exp[17]}

-----------------------------------
"""

    prompt = f"""
You are an AI research assistant for a chemistry laboratory.

Generate a professional {report_type} based ONLY on the
experimental records provided below.

The report is intended for a researcher and their supervisor.

Use the following structure:

# Research Summary

Provide a concise overview of the main research activities
during the reporting period.

## Experiments Conducted

Summarize what experiments were performed and why.

## Key Findings

Identify important experimental results, trends and comparisons.

## Challenges and Unresolved Issues

Identify failed experiments, partial successes, low yields,
incomplete conversion or other problems recorded in the data.

## Next Steps

Summarize the planned future experiments based on the recorded
next-step information.

## Key Metrics

Include:
- Number of experiments
- Average recorded yield
- Number of successful experiments
- Number of partial successes
- Number of failed experiments

Important rules:

- Use ONLY information contained in the experimental records.
- Do NOT invent experimental results.
- Do NOT invent mechanisms.
- Do NOT invent scientific conclusions.
- Clearly distinguish observation from interpretation.
- If information is missing, say so rather than guessing.
- Use professional scientific English.
- Keep the report concise and suitable for a research group meeting.

Experimental records:

{experiment_text}
"""

    return cloud_ai_chat(
        prompt
    )

def generate_professional_procedure(
    experiment_name,
    starting_material,
    sm_mass,
    sm_mmol,
    reagents,
    solvent,
    solvent_volume,
    temperature,
    reaction_time,
    observation
):

    reagent_text = ""

    for reagent in reagents:

        chemical = reagent.get("Chemical", "")

        if not chemical:
            continue

        reagent_text += (
            f"{chemical}: "
            f"{reagent.get('Required mass (mg)', 0):.2f} mg, "
            f"{reagent.get('Required mmol', 0):.3f} mmol, "
            f"{reagent.get('Equiv', 0):.3f} equiv; "
        )

    prompt = f"""
You are a chemistry research assistant.

Convert the following structured experimental information into
a concise, professional experimental procedure suitable for
a laboratory notebook.

Experiment:
{experiment_name}

Starting material:
{starting_material}
Mass: {sm_mass:.2f} mg
Amount: {sm_mmol:.3f} mmol
Equivalent: 1.00 equiv

Reagents:
{reagent_text}

Solvent:
{solvent}
Volume: {solvent_volume:.2f} mL

Temperature:
{temperature:.1f} °C

Reaction time:
{reaction_time:.1f} h

Raw experimental observations:
{observation}

Requirements:

- Use professional scientific English.
- Write in past tense.
- Preserve all numerical values exactly.
- Do not invent reagents, quantities, work-up steps,
  purification methods or analytical results.
- Only include information provided above.
- If a procedure detail is missing, omit it rather than guessing.
- Write one concise experimental paragraph.
"""

    return cloud_ai_chat(
        prompt
    )

# --------------------------------------------------
# PAGE SETTINGS
# --------------------------------------------------

st.set_page_config(
    page_title="LabFlow AI",
    page_icon="🧪",
    layout="wide"
)

# --------------------------------------------------
# SUPABASE CONNECTION
# --------------------------------------------------

@st.cache_resource
def init_supabase():
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)

supabase: Client = init_supabase()

# --------------------------------------------------
# AUTH CLIENT
# --------------------------------------------------

def get_auth_client():

    if "auth_client" not in st.session_state:

        st.session_state.auth_client = create_client(
            st.secrets["SUPABASE_URL"],
            st.secrets["SUPABASE_KEY"]
        )

    return st.session_state.auth_client


auth_client = get_auth_client()

# --------------------------------------------------
# CLOUD AI CONNECTION
# --------------------------------------------------

@st.cache_resource
def init_gemini():

    return genai.Client(
        api_key=st.secrets["GEMINI_API_KEY"]
    )


gemini_client = init_gemini()


def cloud_ai_chat(prompt):

    response = gemini_client.models.generate_content(
        model="gemini-3.6-flash",
        contents=prompt
    )

    return response.text

try:
    test_response = (
        supabase
        .table("experiments")
        .select("id")
        .limit(1)
        .execute()
    )

    print("SUPABASE CONNECTION SUCCESS")

except Exception as e:

    print(
        "SUPABASE CONNECTION FAILED:",
        e
    )

st.markdown(
    """
    <style>

    /* --------------------------------------------------
       GLOBAL
    -------------------------------------------------- */

    .block-container {
        padding-top: 2rem;
        padding-bottom: 3rem;
        max-width: 1400px;
    }

    h1, h2, h3 {
        letter-spacing: -0.02em;
    }

    /* --------------------------------------------------
       SIDEBAR
    -------------------------------------------------- */

    section[data-testid="stSidebar"] {
        border-right: 1px solid rgba(128, 128, 128, 0.18);
    }

    section[data-testid="stSidebar"] .block-container {
        padding-top: 1.5rem;
    }

    /* --------------------------------------------------
       METRIC CARDS
    -------------------------------------------------- */

    div[data-testid="stMetric"] {
        border: 1px solid rgba(128, 128, 128, 0.20);
        border-radius: 14px;
        padding: 16px;
        background: rgba(128, 128, 128, 0.04);
    }

    div[data-testid="stMetricLabel"] {
        font-size: 0.9rem;
        font-weight: 600;
    }

    div[data-testid="stMetricValue"] {
        font-size: 1.7rem;
        font-weight: 700;
    }

    /* --------------------------------------------------
       BUTTONS
    -------------------------------------------------- */

    .stButton > button,
    .stDownloadButton > button {
        border-radius: 10px;
        font-weight: 600;
        min-height: 42px;
    }

    /* --------------------------------------------------
       INPUTS
    -------------------------------------------------- */

    div[data-baseweb="input"] > div,
    div[data-baseweb="select"] > div,
    textarea {
        border-radius: 10px !important;
    }

    /* --------------------------------------------------
       TABLE / DATAFRAME
    -------------------------------------------------- */

    div[data-testid="stDataFrame"] {
        border-radius: 12px;
        overflow: hidden;
    }

    /* --------------------------------------------------
       EXPANDERS
    -------------------------------------------------- */

    details {
        border-radius: 10px !important;
    }

    /* --------------------------------------------------
       HIDE STREAMLIT DEFAULT UI
    -------------------------------------------------- */

    #MainMenu {
        visibility: hidden;
    }

    footer {
        visibility: hidden;
    }

    header[data-testid="stHeader"] {
        background: transparent;
    }

    </style>
    """,
    unsafe_allow_html=True
)

# --------------------------------------------------
# AUTHENTICATION
# --------------------------------------------------

if "user" not in st.session_state:
    st.session_state.user = None


if st.session_state.user is None:

    st.markdown(
        """
        <div style="text-align:center; margin-top:70px;">
            <h1 style="font-size:2.8rem; margin-bottom:8px;">
                🧪 LabFlow AI
            </h1>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown(
        "<p style='text-align:center; font-size:1.1rem; opacity:0.65;'>"
        "From experiments to intelligence."
        "</p>",
        unsafe_allow_html=True
    )
    login_tab, signup_tab = st.tabs(
        [
            "Log In",
            "Sign Up"
        ]
    )


    # --------------------------------------------------
    # LOGIN
    # --------------------------------------------------

    with login_tab:

        login_email = st.text_input(
            "Email",
            key="login_email"
        )

        login_password = st.text_input(
            "Password",
            type="password",
            key="login_password"
        )

        if st.button(
            "Log In",
            type="primary",
            use_container_width=True
        ):

            try:

                login_response = (
                    auth_client
                    .auth
                    .sign_in_with_password({
                        "email": login_email,
                        "password": login_password
                    })
                )

                if login_response.user:

                    st.session_state.user = (
                        login_response.user
                    )

                    st.success(
                        "Login successful."
                    )

                    st.rerun()

            except Exception as e:

                st.error(
                    f"Login failed: {e}"
                )


    # --------------------------------------------------
    # SIGN UP
    # --------------------------------------------------

    with signup_tab:

        signup_email = st.text_input(
            "Email",
            key="signup_email"
        )

        signup_password = st.text_input(
            "Password",
            type="password",
            key="signup_password"
        )

        signup_password_confirm = st.text_input(
            "Confirm Password",
            type="password",
            key="signup_password_confirm"
        )

        if st.button(
            "Create Account",
            use_container_width=True
        ):

            if not signup_email:

                st.warning(
                    "Please enter an email address."
                )

            elif len(signup_password) < 6:

                st.warning(
                    "Password must contain at least 6 characters."
                )

            elif (
                signup_password
                != signup_password_confirm
            ):

                st.warning(
                    "Passwords do not match."
                )

            else:

                try:

                    signup_response = (
                        auth_client
                        .auth
                        .sign_up({
                            "email": signup_email,
                            "password": signup_password
                        })
                    )

                    if signup_response.user:

                        st.success(
                            "Account created successfully. "
                            "You can now log in."
                        )

                except Exception as e:

                    st.error(
                        f"Sign up failed: {e}"
                    )


    st.stop()

# --------------------------------------------------
# SIDEBAR NAVIGATION
# --------------------------------------------------

sidebar_brand = """
<div style="padding:6px 4px 18px 4px;">
<div style="font-size:1.45rem; font-weight:750; letter-spacing:-0.03em;">🧪 LabFlow AI</div>
<div style="margin-top:4px; font-size:0.85rem; opacity:0.65;">From experiments to intelligence.</div>
</div>
"""

st.sidebar.markdown(
    sidebar_brand,
    unsafe_allow_html=True
)

# --------------------------------------------------
# CURRENT USER
# --------------------------------------------------

current_user = st.session_state.user

if current_user:

    st.sidebar.caption(
        "Signed in as"
    )

    st.sidebar.markdown(
        f"**{current_user.email}**"
    )

    if st.sidebar.button(
        "🚪 Log Out",
        use_container_width=True
    ):

        try:

            auth_client.auth.sign_out()

        except Exception:
            pass

        st.session_state.user = None

        if "auth_client" in st.session_state:
            del st.session_state.auth_client

        st.rerun()

st.sidebar.divider()

role = st.sidebar.selectbox(
    "View As",
    [
        "Researcher",
        "Lab Manager"
    ]
)

if role == "Researcher":

    navigation_options = [
        "🏠 Home",
        "🧪 New Experiment",
        "📚 Experiment History",
        "📝 Lab Note Generator",
        "📊 Dashboard",
        "✨ AI Reports",
        "🧠 Research Memory"
    ]

else:

    navigation_options = [
        "🏠 Home",
        "📊 Dashboard",
        "📦 Inventory",
        "⚙️ Lab Manager",
        "✨ AI Reports",
        "🧠 Research Memory"
    ]

page = st.sidebar.radio(
    "Workspace",
    navigation_options
)

st.sidebar.divider()

st.sidebar.info(
    "LabFlow AI helps researchers record experiments, "
    "track progress and transform laboratory data into insights."
)

if page == "🏠 Home":

    hero_html = """
<div style="padding:18px 0 12px 0;">
<div style="font-size:3rem; font-weight:800; letter-spacing:-0.05em; line-height:1.05;">LabFlow AI</div>
<div style="font-size:1.45rem; font-weight:500; margin-top:10px; opacity:0.75;">From experiments to intelligence.</div>
<div style="max-width:780px; margin-top:14px; font-size:1.03rem; line-height:1.7; opacity:0.72;">An AI-powered chemistry laboratory platform for experimental recording, research analytics, automated reporting and laboratory management.</div>
</div>
"""

    st.markdown(
        hero_html,
        unsafe_allow_html=True
    )

    st.divider()

    st.caption(
        f"Current workspace: {role}"
    )
    # --------------------------------------------------
    # KEY METRICS
    # --------------------------------------------------

    experiment_response = (
        supabase
        .table("experiments")
        .select(
            "id, yield_percent"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .execute()
    )

    experiment_rows = (
        experiment_response.data
        or []
    )

    total_experiments = len(
        experiment_rows
    )

    yield_values = [
        row["yield_percent"]
        for row in experiment_rows
        if row["yield_percent"] is not None
    ]

    average_yield = (
        sum(yield_values)
        / len(yield_values)
        if yield_values
        else None
    )

    inventory_response = (
        supabase
        .table("inventory")
        .select(
            "id, current_stock, minimum_stock"
        )
        .execute()
    )

    inventory_rows = (
        inventory_response.data
        or []
    )

    total_inventory_items = len(
        inventory_rows
    )

    low_stock_items = len([
        item
        for item in inventory_rows
        if (
            (item["current_stock"] or 0)
            <=
            (item["minimum_stock"] or 0)
        )
    ])   
    col1, col2, col3, col4 = st.columns(4)

    with col1:

        st.metric(
            "Experiments Recorded",
            total_experiments
        )

    with col2:

        st.metric(
            "Average Yield",
            (
                f"{average_yield:.1f}%"
                if average_yield is not None
                else "N/A"
            )
        )

    with col3:

        st.metric(
            "Inventory Items",
            total_inventory_items
        )

    with col4:

        st.metric(
            "Low Stock Alerts",
            low_stock_items
        )

    st.divider()

    st.markdown("## 🎬 Demo Mode")

    st.caption(
        "Load a complete demonstration dataset for presentations and testing."
    )

    col1, col2 = st.columns(2)

    with col1:

         if st.button(
            "🎬 Load Demo Data",
            use_container_width=True
        ):

            if demo_data_exists():

                st.warning(
                    "Demo data is already loaded."
                )

            else:

                load_demo_data()

                st.success(
                    "Demo dataset loaded successfully."
                )

                st.rerun()

    with col2:

        if st.button(
            "🧹 Clear Demo Data",
            use_container_width=True
        ):

            clear_demo_data()

            st.success(
                "Demo dataset removed."
            )

            st.rerun()

    st.divider()

    # --------------------------------------------------
    # WORKFLOW
    # --------------------------------------------------

    st.markdown(
        "## One workflow. Complete laboratory intelligence."
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        feature_card(
            "🧪",
            "Record",
            "Capture reagents, quantities, reaction conditions, "
            "observations and research context directly during experiments."
        )

        feature_card(
            "📊",
            "Analyse",
            "Track yield trends, experimental activity, solvent usage "
            "and research performance automatically."
        )

    with col2:

        feature_card(
            "✨",
            "Generate",
            "Convert structured data into professional lab notes, "
            "AI procedures and weekly or monthly reports."
        )

        feature_card(
            "🧠",
            "Remember",
            "Search experimental history in natural language and "
            "compare previous results using laboratory memory."
        )

    with col3:

        feature_card(
            "📦",
            "Manage",
            "Track inventory consumption automatically and receive "
            "low-stock alerts before critical materials run out."
        )

        feature_card(
            "💰",
            "Control",
            "Monitor experiment costs, solvent consumption and laboratory "
            "spending from one management dashboard."
        )
    st.divider()

    # --------------------------------------------------
    # CORE VALUE
    # --------------------------------------------------

    st.markdown("## Why LabFlow?")

    st.info(
        "Record experimental information once. "
        "LabFlow automatically transforms the same data "
        "into calculations, lab records, dashboards, "
        "reports, inventory updates and laboratory insights."
    )

    st.divider()

    st.markdown("## LabFlow Workflow")

    st.code(
        """
Experiment
    ↓
Smart Data Entry
    ↓
Stoichiometry & Cost Calculation
    ↓
Structured Lab Database
    ↓
 ┌──────────────┬──────────────┬──────────────┐
 ↓              ↓              ↓
Lab Note      Dashboard      Inventory
 ↓              ↓              ↓
AI Procedure  AI Reports    Cost Control
      \\          |          /
       \\         |         /
          Research Memory
        """,
        language=None
    )

if page == "🧪 New Experiment":

    st.header("🧪 New Experiment")

    st.caption(
        "Record experimental conditions, reagents, observations and results."
    )

    st.divider()

    col1, col2, col3 = st.columns(3)

    with col1:
        experiment_name = st.text_input(
            "Experiment Name",
            placeholder="e.g. Suzuki Coupling"
        )

    with col2:
        researcher = st.text_input(
            "Researcher",
            placeholder="e.g. Alex Chen"
        )

    with col3:
        experiment_date = st.date_input(
            "Date",
            value=date.today()
        )

    st.divider()

    # --------------------------------------------------
    # REACTION SETUP
    # --------------------------------------------------

    st.header("Reaction Setup")

    st.subheader("Starting Material")

    col1, col2, col3 = st.columns(3)

    with col1:
        sm_name = st.text_input(
            "Chemical Name",
            placeholder="e.g. 4-Bromoacetophenone"
        )

    with col2:
        sm_mw = st.number_input(
            "Molecular Weight (g/mol)",
            min_value=0.0,
            value=0.0,
            step=0.01
        )

    with col3:
        sm_mass = st.number_input(
            "Mass (mg)",
            min_value=0.0,
            value=0.0,
            step=1.0
        )

    # Calculate mmol
    if sm_mw > 0:
        sm_mmol = sm_mass / sm_mw
    else:
        sm_mmol = 0.0

    col1, col2 = st.columns(2)

    with col1:
        st.metric(
            "Amount",
            f"{sm_mmol:.3f} mmol"
        )

    with col2:
        st.metric(
            "Equivalent",
            "1.00 equiv"
        )

    st.divider()

    # --------------------------------------------------
    # REAGENT TABLE
    # --------------------------------------------------

    st.subheader("Reagents")

    st.caption(
        "Add reagents, catalysts, bases and additives. "
        "Required amounts are calculated automatically "
        "from the starting material."
    )

    # Create reagent storage
    if "reagents" not in st.session_state:
        st.session_state.reagents = [
            {
                "Chemical": "",
                "Role": "Reagent",
                "MW (g/mol)": 0.0,
                "Equiv": 1.0,
                "Cost (HKD/g)": 0.0
            }
        ]

    # Editable reagent table
    edited_reagents = st.data_editor(
        st.session_state.reagents,
        num_rows="dynamic",
        use_container_width=True,
        column_config={
            "Chemical": st.column_config.TextColumn(
                "Chemical"
            ),
            "Role": st.column_config.SelectboxColumn(
                "Role",
                options=[
                    "Reagent",
                    "Base",
                    "Catalyst",
                    "Ligand",
                    "Additive"
                ]
            ),
            "MW (g/mol)": st.column_config.NumberColumn(
                "MW (g/mol)",
                min_value=0.0,
                format="%.2f"
            ),
            "Equiv": st.column_config.NumberColumn(
                "Equiv",
                min_value=0.0,
                format="%.3f"
            ),
            "Cost (HKD/g)": st.column_config.NumberColumn(
                "Cost (HKD/g)",
                min_value=0.0,
                format="%.2f"
            )
        },
        key="reagent_editor"
    )

    st.session_state.reagents = edited_reagents

    # Calculate reagent quantities
    calculated_reagents = []

    for reagent in edited_reagents:

        mw = reagent.get("MW (g/mol)", 0.0) or 0.0
        equiv = reagent.get("Equiv", 0.0) or 0.0

        mmol = sm_mmol * equiv
        mass_mg = mmol * mw
        cost_per_g = reagent.get("Cost (HKD/g)", 0.0) or 0.0

        reagent_cost = (
            mass_mg / 1000
        ) * cost_per_g

        calculated_reagents.append({
            "Chemical": reagent.get("Chemical", "") or "",
            "Role": reagent.get("Role", "Reagent") or "Reagent",
            "MW (g/mol)": mw,
            "Equiv": equiv,
            "Required mmol": round(mmol, 3),
            "Required mass (mg)": round(mass_mg, 1),
            "Estimated Cost (HKD)": round(reagent_cost, 2)
        })

    st.markdown("#### Calculated Quantities")

    st.dataframe(
        calculated_reagents,
        use_container_width=True,
        hide_index=True
    )

    st.divider()

    st.header("Reaction Conditions")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        solvent = st.text_input(
            "Solvent",
            placeholder="e.g. THF"
        )

    with col2:
        solvent_volume = st.number_input(
            "Volume (mL)",
            min_value=0.0,
            value=0.0,
            step=0.1
        )

    with col3:
        solvent_cost_per_l = st.number_input(
            "Solvent Cost (HKD/L)",
            min_value=0.0,
            value=0.0,
            step=1.0
        )

    with col4:
        temperature = st.number_input(
            "Temperature (°C)",
            value=25.0,
            step=5.0
        )

    reaction_time = st.number_input(
        "Reaction Time (hours)",
        min_value=0.0,
        value=1.0,
        step=0.5
    )

    solvent_cost = (
        solvent_volume / 1000
    ) * solvent_cost_per_l

    st.divider()

    # --------------------------------------------------
    # OBSERVATIONS
    # --------------------------------------------------

    st.header("Observations")

    observation = st.text_area(
        "Experimental Observations",
        placeholder=(
            "e.g. Reaction mixture turned dark brown after "
            "20 minutes. TLC showed remaining starting material..."
        ),
        height=120
    )

    yield_percent = st.number_input(
        "Isolated Yield (%)",
        min_value=0.0,
        max_value=100.0,
        value=0.0,
        step=1.0
    )

    st.divider()

    # --------------------------------------------------
    # RESEARCH CONTEXT
    # --------------------------------------------------

    st.header("Research Context")

    objective = st.text_area(
        "🎯 Objective",
        placeholder=(
            "What is the purpose of this experiment? "
            "e.g. Investigate whether increasing catalyst "
            "loading improves reaction conversion."
        ),
        height=100
    )

    status = st.selectbox(
        "📌 Experiment Status",
        [
            "In Progress",
            "Successful",
            "Partial Success",
            "Failed"
        ]
    )

    key_result = st.text_area(
        "🔬 Key Result",
        placeholder=(
            "What was the most important result? "
            "e.g. Increasing catalyst loading to 5 mol% "
            "improved the isolated yield to 70%."
        ),
        height=100
    )

    next_step = st.text_area(
        "➡️ Next Step",
        placeholder=(
            "What should be done next? "
            "e.g. Test higher catalyst loading and "
            "increase the reaction temperature."
        ),
        height=100
    )

    st.divider()

    # --------------------------------------------------
    # CALCULATE TOTAL EXPERIMENT COST
    # --------------------------------------------------

    total_reagent_cost = sum(
        reagent["Estimated Cost (HKD)"]
        for reagent in calculated_reagents
    )

    estimated_total_cost = (
        total_reagent_cost
        + solvent_cost
    )
    st.metric(
        "Estimated Experiment Cost",
        f"HK${estimated_total_cost:.2f}"
    )    

    # --------------------------------------------------
    # CREATE EXPERIMENT
    # --------------------------------------------------

    if "last_saved_signature" not in st.session_state:
        st.session_state.last_saved_signature = None

    if st.button(
        "Create Experiment",
        type="primary",
        use_container_width=True
    ):

        if not experiment_name:

            st.error(
                "Please enter an experiment name."
            )

        elif not sm_name:

            st.error(
                "Please enter a starting material."
            )

        else:

            inventory_shortages = check_experiment_inventory(
                calculated_reagents,
                solvent,
                solvent_volume
            )

            if inventory_shortages:

                st.error(
                    "Insufficient inventory. "
                    "The experiment was not created."
                )

                for shortage in inventory_shortages:

                    st.warning(
                        f"{shortage['chemical']}: "
                        f"required {shortage['required']:.3f} "
                        f"{shortage['unit']}, "
                        f"available {shortage['available']:.3f} "
                        f"{shortage['unit']}."
                    )

            else:

                save_signature = (
                    experiment_name,
                    str(experiment_date),
                    sm_name,
                    round(sm_mass, 4),
                    round(yield_percent, 2)
                )

                if (
                    st.session_state.last_saved_signature
                    == save_signature
                ):

                    st.warning(
                        "This experiment appears to have "
                        "already been saved."
                    )

                else:

                    experiment_id = save_experiment(
                        experiment_name,
                        researcher,
                        experiment_date,
                        sm_name,
                        sm_mw,
                        sm_mass,
                        sm_mmol,
                        calculated_reagents,
                        solvent,
                        solvent_volume,
                        temperature,
                        reaction_time,
                        yield_percent,
                        observation,
                        objective,
                        status,
                        key_result,
                        next_step,
                        solvent_cost_per_l,
                        estimated_total_cost
                    )
                    deduct_experiment_inventory(
                        calculated_reagents,
                        solvent,
                        solvent_volume
                    )

                    mark_inventory_deducted(
                        experiment_id
                    )

                    st.session_state.last_saved_signature = (
                        save_signature
                    )

                    st.success(
                        f"Experiment EXP-{experiment_id:04d} "
                        f"saved successfully!"
                    )

                    st.subheader(
                        "Experiment Summary"
                    )

                    st.write(
                        f"**Experiment:** {experiment_name}"
                    )

                    st.write(
                        f"**Researcher:** {researcher}"
                    )

                    st.write(
                        f"**Date:** {experiment_date}"
                    )

                    st.write("---")

                    st.write(
                        f"**Starting Material:** "
                        f"{sm_name} — "
                        f"{sm_mass:.1f} mg, "
                        f"{sm_mmol:.3f} mmol, "
                        f"1.00 equiv"
                    )

                    for reagent in calculated_reagents:

                        if reagent["Chemical"]:

                            st.write(
                                f"**{reagent['Role']}:** "
                                f"{reagent['Chemical']} — "
                                f"{reagent['Required mass (mg)']:.1f} mg, "
                                f"{reagent['Required mmol']:.3f} mmol, "
                                f"{reagent['Equiv']:.3f} equiv"
                            )

                    st.write(
                        f"**Solvent:** {solvent} "
                        f"({solvent_volume:.1f} mL)"
                    )

                    st.write(
                        f"**Conditions:** "
                        f"{temperature:.0f} °C, "
                        f"{reaction_time:.1f} h"
                    )

                    st.write(
                        f"**Yield:** "
                        f"{yield_percent:.1f}%"
                    )

                    if observation:

                        st.write(
                            f"**Observations:** "
                            f"{observation}"
                        )

                    st.write("---")

                    st.write(
                        f"**Objective:** {objective}"
                    )

                    st.write(
                        f"**Status:** {status}"
                    )

                    st.write(
                        f"**Key Result:** {key_result}"
                    )

                    st.write(
                        f"**Next Step:** {next_step}"
                    )   

# --------------------------------------------------
# EXPERIMENT HISTORY PAGE
# --------------------------------------------------

if page == "📚 Experiment History":

    st.header("📚 Experiment History")
    st.caption(
        "Browse all experiments saved in the LabFlow database."
    )
    st.divider()

    response = (
        supabase
        .table("experiments")
        .select(
            "id, "
            "experiment_name, "
            "researcher, "
            "experiment_date, "
            "starting_material, "
            "yield_percent, "
            "status"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .order(
            "id",
            desc=True
        )
        .execute()
    )

    experiments = response.data    

    if experiments:

        history_data = []

        for exp in experiments:

            history_data.append({
                "Experiment ID": (
                    f"EXP-{exp['id']:04d}"
                ),
                "Experiment": exp[
                    "experiment_name"
                ],
                "Researcher": exp[
                    "researcher"
                ],
                "Date": exp[
                    "experiment_date"
                ],
                "Starting Material": exp[
                    "starting_material"
                ],
                "Yield (%)": exp[
                    "yield_percent"
                ],
                "Status": (
                    exp["status"]
                    or "Not specified"
                )
            })

        st.dataframe(
            history_data,
            use_container_width=True,
            hide_index=True
        )
        st.divider()

        st.subheader("🗑️ Delete Experiment")

        st.caption(
            "If this experiment previously deducted inventory, "
            "the consumed reagents and solvent will be restored automatically."
        )
        
        delete_options = {
            (
                f"EXP-{exp['id']:04d} | "
                f"{exp['experiment_name']} | "
                f"{exp['experiment_date']}"
            ): exp["id"]
            for exp in experiments
        }

        selected_delete_label = st.selectbox(
            "Select Experiment to Delete",
            list(delete_options.keys())
        )

        selected_delete_id = delete_options[
            selected_delete_label
        ]

        confirm_delete = st.checkbox(
            "I understand this action cannot be undone.",
            key=f"confirm_delete_{selected_delete_id}"
        )

        if st.button(
            "🗑️ Delete Selected Experiment",
            use_container_width=True
        ):

            if not confirm_delete:

                st.warning(
                    "Please confirm before deleting the experiment."
                )

            else:

                deleted = delete_experiment(
                    selected_delete_id
                )

                if deleted:

                    st.success(
                        f"EXP-{selected_delete_id:04d} "
                        f"deleted successfully."
                    )

                    st.rerun()

                else:

                    st.error(
                        "Experiment could not be found."
                    )
    else:
        st.info("No experiments saved yet.")

if page == "📝 Lab Note Generator":

    st.header("📝 Lab Note Generator")

    st.caption(
        "Generate a structured laboratory record from saved experimental data."
    )
    st.divider()

    response = (
        supabase
        .table("experiments")
        .select(
            "id, experiment_name, experiment_date"
        )
        .eq(
            "user_id",
            st.session_state.user.id
        )
        .order(
            "id",
            desc=True
        )
        .execute()
    )

    experiment_options = [
        (
            exp["id"],
            exp["experiment_name"],
            exp["experiment_date"]
        )
        for exp in response.data
    ]

    if not experiment_options:

        st.info(
            "No experiments are available. "
            "Create an experiment first."
        )

    else:

        experiment_labels = {
            (
                f"EXP-{exp[0]:04d} | "
                f"{exp[1]} | "
                f"{exp[2]}"
            ): exp[0]
            for exp in experiment_options
        }

        selected_label = st.selectbox(
            "Select Experiment",
            list(experiment_labels.keys())
        )

        selected_id = experiment_labels[
            selected_label
        ]

        experiment = get_experiment_by_id(
            selected_id
        )

        if experiment:

            (
                experiment_id,
                experiment_name,
                researcher,
                experiment_date,
                starting_material,
                sm_mw,
                sm_mass,
                sm_mmol,
                reagents_json,
                solvent,
                solvent_volume,
                temperature,
                reaction_time,
                yield_percent,
                observation,
                objective,
                status,
                key_result,
                next_step,
                estimated_total_cost,
                saved_ai_procedure
            ) = experiment

            try:
                reagents = json.loads(
                    reagents_json
                )

            except:
                reagents = []

            st.divider()

            st.subheader(
                f"EXP-{experiment_id:04d} — "
                f"{experiment_name}"
            )

            st.write(
                f"**Researcher:** {researcher}"
            )

            st.write(
                f"**Date:** {experiment_date}"
            )

            st.write(
                f"**Status:** "
                f"{status or 'Not specified'}"
            )

            st.divider()

            # ------------------------------------------
            # OBJECTIVE
            # ------------------------------------------

            st.markdown("### 🎯 Objective")

            st.write(
                objective
                or "No objective recorded."
            )

            # ------------------------------------------
            # REACTION SETUP
            # ------------------------------------------

            st.markdown("### 🧪 Reaction Setup")

            st.markdown(
                "#### Starting Material"
            )

            st.write(
                f"**{starting_material}**"
            )

            st.write(
                f"{sm_mass:.2f} mg | "
                f"{sm_mmol:.3f} mmol | "
                f"1.00 equiv | "
                f"MW {sm_mw:.2f} g/mol"
            )

            if reagents:

                st.markdown("#### Reagents")

                for reagent in reagents:

                    chemical = reagent.get(
                        "Chemical",
                        ""
                    )

                    if not chemical:
                        continue

                    role = reagent.get(
                        "Role",
                        "Reagent"
                    )

                    mass = reagent.get(
                        "Required mass (mg)",
                        0
                    )

                    mmol = reagent.get(
                        "Required mmol",
                        0
                    )

                    equiv = reagent.get(
                        "Equiv",
                        0
                    )

                    st.write(
                        f"**{chemical}** "
                        f"({role}) — "
                        f"{mass:.2f} mg | "
                        f"{mmol:.3f} mmol | "
                        f"{equiv:.3f} equiv"
                    )

            # ------------------------------------------
            # REACTION CONDITIONS
            # ------------------------------------------

            st.markdown(
                "### 🌡️ Reaction Conditions"
            )

            st.write(
                f"**Solvent:** "
                f"{solvent or 'Not recorded'} "
                f"({solvent_volume or 0:.2f} mL)"
            )

            st.write(
                f"**Temperature:** "
                f"{temperature or 0:.1f} °C"
            )

            st.write(
                f"**Reaction Time:** "
                f"{reaction_time or 0:.1f} h"
            )

            # ------------------------------------------
            # OBSERVATIONS
            # ------------------------------------------

            st.markdown(
                "### 👁️ Experimental Observations"
            )

            st.write(
                observation
                or "No observations recorded."
            )

            st.divider()

            st.markdown(
                "### ✨ AI Experimental Procedure"
            )

            if st.button(
                "✨ Generate Professional Procedure",
                use_container_width=True
            ):

                with st.spinner(
                    "Generating experimental procedure..."
                ):

                    try:

                        ai_procedure = (
                            generate_professional_procedure(
                                experiment_name,
                                starting_material,
                                sm_mass,
                                sm_mmol,
                                reagents,
                                solvent,
                                solvent_volume,
                                temperature,
                                reaction_time,
                                observation
                            )
                        )

                        save_ai_procedure(
                            experiment_id,
                            ai_procedure
                        )

                        st.session_state[
                            f"ai_procedure_{experiment_id}"
                        ] = ai_procedure

                    except Exception as e:

                        st.error(
                            f"Procedure generation failed: {e}"
                        )

            procedure_key = (
                f"ai_procedure_{experiment_id}"
            )

            current_ai_procedure = (
                st.session_state.get(
                    procedure_key
                )
                or saved_ai_procedure
            )

            if current_ai_procedure:

                st.success(
                    "Professional procedure available."
                )

                st.write(
                    current_ai_procedure
                )


            # ------------------------------------------
            # RESULTS
            # ------------------------------------------

            st.markdown("### 📊 Result")

            if yield_percent is not None:

                st.write(
                    f"**Isolated Yield:** "
                    f"{yield_percent:.1f}%"
                )

            else:

                st.write(
                    "**Isolated Yield:** "
                    "Not recorded"
                )

            st.write(
                f"**Experiment Status:** "
                f"{status or 'Not specified'}"
            )

            # ------------------------------------------
            # KEY RESULT
            # ------------------------------------------

            st.markdown("### 🔬 Key Result")

            st.write(
                key_result
                or "No key result recorded."
            )

            # ------------------------------------------
            # NEXT STEP
            # ------------------------------------------

            st.markdown("### ➡️ Next Step")

            st.write(
                next_step
                or "No next step recorded."
            )

            st.divider()

            st.metric(
                "Estimated Experiment Cost",
                f"HK${estimated_total_cost:.2f}"
                if estimated_total_cost is not None
                else "Not recorded"
            )
        
            # ------------------------------------------
            # EXPERIMENT ATTACHMENTS
            # ------------------------------------------

            st.divider()

            st.markdown(
                "### 📎 Experimental Attachments"
            )

            st.caption(
                "Upload TLC images, reaction photos, "
                "NMR spectra, analytical data and other "
                "experimental files."
            )

            attachment_category = st.selectbox(
                "Attachment Type",
                [
                    "TLC",
                    "Reaction Photo",
                    "NMR",
                    "IR",
                    "UV-Vis",
                    "HPLC",
                    "GC-MS",
                    "LC-MS",
                    "XRD",
                    "SEM",
                    "Raw Data",
                    "Other"
                ],
                key=f"attachment_category_{experiment_id}"
            )

            uploaded_attachment = st.file_uploader(
                "Choose a file",
                type=[
                    "png",
                    "jpg",
                    "jpeg",
                    "pdf",
                    "csv",
                    "txt",
                    "xlsx"
                ],
                key=f"attachment_upload_{experiment_id}"
            )

            attachment_notes = st.text_area(
                "Attachment Notes",
                placeholder=(
                    "e.g. TLC after 2 h; "
                    "eluent = PE/EtOAc 4:1"
                ),
                key=f"attachment_notes_{experiment_id}"
            )

            if st.button(
                "📎 Save Attachment",
                use_container_width=True,
                key=f"save_attachment_{experiment_id}"
            ):

                if uploaded_attachment is None:

                    st.warning(
                        "Please select a file first."
                    )

                else:

                    save_attachment(
                        experiment_id,
                        uploaded_attachment,
                        attachment_category,
                        attachment_notes
                    )

                    st.success(
                        f"{uploaded_attachment.name} "
                        f"saved to EXP-{experiment_id:04d}."
                    )

                    st.rerun()

            st.markdown("#### Saved Attachments")

            saved_attachments = get_attachments(
                experiment_id
            )

            if not saved_attachments:

                st.info(
                    "No attachments saved for this experiment."
                )

            else:

                st.write(
                    f"**{len(saved_attachments)} file(s) attached**"
                )

                for attachment in saved_attachments:

                    attachment_id = attachment[0]
                    file_name = attachment[1]
                    file_type = attachment[2]
                    category = attachment[3]
                    file_data = attachment[4]
                    notes = attachment[5]
                    uploaded_at = attachment[6]

                    with st.expander(
                        f"📎 {category} — {file_name}"
                    ):

                        st.write(
                            f"**Category:** {category}"
                        )

                        st.write(
                            f"**Uploaded:** {uploaded_at}"
                        )

                        if notes:

                            st.write(
                                f"**Notes:** {notes}"
                            )

                        # ----------------------------------
                        # IMAGE PREVIEW
                        # ----------------------------------

                        if file_type in [
                            "image/png",
                            "image/jpeg"
                        ]:

                            st.image(
                                file_data,
                                caption=file_name,
                                use_container_width=True
                            )

                        # ----------------------------------
                        # DOWNLOAD
                        # ----------------------------------

                        st.download_button(
                            label=f"⬇ Download {file_name}",
                            data=file_data,
                            file_name=file_name,
                            mime=(
                                file_type
                                or "application/octet-stream"
                            ),
                            key=(
                                f"download_attachment_"
                                f"{attachment_id}"
                            ),
                            use_container_width=True
                        )


            word_file = generate_word_lab_note(
                experiment,
                current_ai_procedure
            )

            pdf_file = generate_pdf_lab_note(
                experiment,
                current_ai_procedure
            )

            st.download_button(
                label="📝 Download Word Lab Note",
                data=word_file,
                file_name=(
                    f"EXP-{experiment_id:04d}_"
                    f"{experiment_name.replace(' ', '_')}_LabNote.docx"
                ),
                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),
                use_container_width=True
            )

            pdf_file = generate_pdf_lab_note(
                experiment,
                current_ai_procedure
            )

            st.download_button(
                label="📄 Download PDF Lab Note",
                data=pdf_file,
                file_name=(
                    f"EXP-{experiment_id:04d}_"
                    f"{experiment_name.replace(' ', '_')}_LabNote.pdf"
                ),
                mime="application/pdf",
                use_container_width=True
            )
            

# --------------------------------------------------
# DASHBOARD PAGE
# --------------------------------------------------
if page == "📊 Dashboard":

    st.header("📊 Research Dashboard")

    st.caption(
        "Track experimental activity, yield trends and laboratory usage."
    )

    st.divider()

    experiments = get_all_experiments()

    if not experiments:
        st.info(
            "No experiment data available yet. "
            "Create some experiments first."
        )

    else:

        # --------------------------------------------------
        # TIME FILTER
        # --------------------------------------------------

        time_filter = st.selectbox(
            "Time Period",
            [
                "All Time",
                "This Week",
                "This Month"
            ]
        )

        today = date.today()

        filtered_experiments = []

        for exp in experiments:

            try:
                exp_date = datetime.strptime(
                    exp[3],
                    "%Y-%m-%d"
                ).date()

            except:
                continue

            if time_filter == "All Time":
                filtered_experiments.append(exp)

            elif time_filter == "This Week":

                start_of_week = today - timedelta(
                    days=today.weekday()
                )

                if exp_date >= start_of_week:
                    filtered_experiments.append(exp)

            elif time_filter == "This Month":

                if (
                    exp_date.year == today.year
                    and exp_date.month == today.month
                ):
                    filtered_experiments.append(exp)

        if not filtered_experiments:

            st.warning(
                f"No experiments found for {time_filter}."
            )

        else:

            # --------------------------------------------------
            # SUMMARY METRICS
            # --------------------------------------------------

            total_experiments = len(
                filtered_experiments
            )

            yields = [
                exp[7]
                for exp in filtered_experiments
                if exp[7] is not None
            ]

            average_yield = (
                sum(yields) / len(yields)
                if yields
                else 0
            )

            successful_experiments = len([
                y for y in yields
                if y >= 50
            ])

            success_rate = (
                successful_experiments
                / len(yields)
                * 100
                if yields
                else 0
            )

            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric(
                    "Experiments",
                    total_experiments
                )

            with col2:
                st.metric(
                    "Average Yield",
                    f"{average_yield:.1f}%"
                )

            with col3:
                st.metric(
                    "Successful",
                    successful_experiments
                )

            with col4:
                st.metric(
                    "Success Rate",
                    f"{success_rate:.1f}%"
                )

            st.divider()

            # --------------------------------------------------
            # EXPERIMENT ACTIVITY
            # --------------------------------------------------

            st.subheader("Experiment Activity")

            activity_data = {}

            for exp in filtered_experiments:

                exp_date = exp[3]

                activity_data[exp_date] = (
                    activity_data.get(exp_date, 0)
                    + 1
                )

            activity_chart = {
                "Date": list(activity_data.keys()),
                "Experiments": list(
                    activity_data.values()
                )
            }

            st.bar_chart(
                activity_chart,
                x="Date",
                y="Experiments"
            )

            st.divider()

            # --------------------------------------------------
            # YIELD TREND
            # --------------------------------------------------

            st.subheader("Yield Trend")

            yield_chart_data = []

            for exp in filtered_experiments:

                if exp[7] is not None:

                    yield_chart_data.append({
                        "Date": exp[3],
                        "Yield (%)": exp[7]
                    })

            if yield_chart_data:

                st.line_chart(
                    yield_chart_data,
                    x="Date",
                    y="Yield (%)"
                )

            else:

                st.info(
                    "No yield data available."
                )

            st.divider()

            # --------------------------------------------------
            # SOLVENT ANALYSIS
            # --------------------------------------------------

            st.subheader("Solvent Usage")

            solvent_usage = {}

            for exp in filtered_experiments:

                solvent = exp[5]
                volume = exp[6] or 0

                if solvent:
                    solvent_usage[solvent] = (
                        solvent_usage.get(
                            solvent,
                            0
                        )
                        + volume
                    )

            if solvent_usage:

                top_solvent = max(
                    solvent_usage,
                    key=solvent_usage.get
                )

                total_solvent = sum(
                    solvent_usage.values()
                )

                col1, col2 = st.columns(2)

                with col1:
                    st.metric(
                        "Most Used Solvent",
                        top_solvent
                    )

                with col2:
                    st.metric(
                        "Total Solvent Used",
                        f"{total_solvent:.1f} mL"
                    )

                solvent_chart_data = {
                    "Solvent":
                        list(solvent_usage.keys()),
                    "Volume (mL)":
                        list(solvent_usage.values())
                }

                st.bar_chart(
                    solvent_chart_data,
                    x="Solvent",
                    y="Volume (mL)"
                )

            else:

                st.info(
                    "No solvent usage data available."
                )

            st.divider()

            # --------------------------------------------------
            # RECENT EXPERIMENTS
            # --------------------------------------------------

            st.subheader("Recent Experiments")

            recent_data = []

            for exp in filtered_experiments[:5]:

                recent_data.append({
                    "Experiment ID":
                        f"EXP-{exp[0]:04d}",
                    "Experiment":
                        exp[1],
                    "Researcher":
                        exp[2],
                    "Date":
                        exp[3],
                    "Starting Material":
                        exp[4],
                    "Solvent":
                        exp[5],
                    "Yield (%)":
                        exp[7]
                })

            st.dataframe(
                recent_data,
                use_container_width=True,
                hide_index=True
            )

# --------------------------------------------------
# AI REPORTS PAGE
# --------------------------------------------------
if page == "✨ AI Reports":

    st.header("✨ AI Research Reports")

    st.caption(
        "Transform structured laboratory records into research summaries."
    )

    st.divider()

    # --------------------------------------------------
    # REPORT TYPE
    # --------------------------------------------------

    report_type = st.selectbox(
        "Report Type",
        [
            "Weekly Report",
            "Monthly Report"
        ]
    )

    today = date.today()

    # --------------------------------------------------
    # GET EXPERIMENTS
    # --------------------------------------------------

    all_report_experiments = get_experiments_for_report()

    selected_experiments = []

    for exp in all_report_experiments:

        try:
            exp_date = datetime.strptime(
                exp[3],
                "%Y-%m-%d"
            ).date()

        except:
            continue

        if report_type == "Weekly Report":

            start_of_week = today - timedelta(
                days=today.weekday()
            )

            end_of_week = start_of_week + timedelta(
                days=6
            )

            if start_of_week <= exp_date <= end_of_week:
                selected_experiments.append(exp)

        elif report_type == "Monthly Report":

            if (
                exp_date.year == today.year
                and exp_date.month == today.month
            ):
                selected_experiments.append(exp)

    # --------------------------------------------------
    # REPORT PERIOD
    # --------------------------------------------------

    if report_type == "Weekly Report":

        start_of_week = today - timedelta(
            days=today.weekday()
        )

        end_of_week = start_of_week + timedelta(
            days=6
        )

        st.info(
            f"Report period: "
            f"{start_of_week.strftime('%d %b %Y')} – "
            f"{end_of_week.strftime('%d %b %Y')}"
        )

    else:

        st.info(
            f"Report period: "
            f"{today.strftime('%B %Y')}"
        )

    # --------------------------------------------------
    # SUMMARY
    # --------------------------------------------------

    st.subheader("Experiments Included")

    if not selected_experiments:

        st.warning(
            "No experiments were found for this period."
        )

    else:

        st.success(
            f"{len(selected_experiments)} experiments "
            f"will be included in this report."
        )

        # --------------------------------------------------
        # SHOW EXPERIMENTS
        # --------------------------------------------------

        for exp in selected_experiments:

            experiment_id = exp[0]
            experiment_name = exp[1]
            researcher = exp[2]
            experiment_date = exp[3]
            starting_material = exp[4]
            yield_percent = exp[12]
            objective = exp[14]
            status = exp[15]
            key_result = exp[16]
            next_step = exp[17]

            with st.expander(
                f"EXP-{experiment_id:04d} — "
                f"{experiment_name}"
            ):

                st.write(
                    f"**Date:** {experiment_date}"
                )

                st.write(
                    f"**Researcher:** {researcher}"
                )

                st.write(
                    f"**Starting Material:** "
                    f"{starting_material}"
                )

                st.write(
                    f"**Yield:** "
                    f"{yield_percent:.1f}%"
                    if yield_percent is not None
                    else "**Yield:** Not recorded"
                )

                st.write(
                    f"**Status:** "
                    f"{status or 'Not specified'}"
                )

                st.write(
                    f"**Objective:** "
                    f"{objective or 'Not recorded'}"
                )

                st.write(
                    f"**Key Result:** "
                    f"{key_result or 'Not recorded'}"
                )

                st.write(
                    f"**Next Step:** "
                    f"{next_step or 'Not recorded'}"
                )
                st.divider()

        if st.button(
            "✨ Generate AI Report",
            type="primary",
            use_container_width=True
        ):

            with st.spinner(
                "Analyzing experimental records..."
            ):

                try:

                    ai_report = generate_ai_report(
                        selected_experiments,
                        report_type
                    )

                    st.session_state[
                        "generated_report"
                    ] = ai_report

                except Exception as e:

                    st.error(
                        f"AI report generation failed: {e}"
                    )

        if "generated_report" in st.session_state:

            st.divider()

            st.subheader("Generated Report")

            st.markdown(
                st.session_state[
                    "generated_report"
                ]
            )

if page == "🧠 Research Memory":

    st.header("🧠 Research Memory")

    st.caption(
        "Search and analyse your experimental history using natural language."
    )

    st.divider()

    experiments_for_search = (
        get_experiments_for_ai_search()
    )

    if not experiments_for_search:

        st.info(
            "No experimental records are available."
        )

    else:

        st.success(
            f"{len(experiments_for_search)} experiments "
            f"available in the database."
        )

        st.markdown("### 🔎 Search Filters")

        # --------------------------------------------------
        # BUILD FILTER OPTIONS
        # --------------------------------------------------

        researcher_options = sorted(
            list({
                exp[2]
                for exp in experiments_for_search
                if exp[2]
            })
        )

        status_options = sorted(
            list({
                exp[15]
                for exp in experiments_for_search
                if exp[15]
            })
        )

        solvent_options = sorted(
            list({
                exp[8]
                for exp in experiments_for_search
                if exp[8]
            })
        )

        # --------------------------------------------------
        # DATE FILTER
        # --------------------------------------------------

        date_mode = st.selectbox(
            "Date Range",
            [
                "All Time",
                "Last 7 Days",
                "Last 30 Days",
                "This Month",
                "Custom"
            ]
        )

        today = date.today()

        start_date_filter = None
        end_date_filter = None

        if date_mode == "Last 7 Days":

            start_date_filter = (
                today - timedelta(days=6)
            )

            end_date_filter = today

        elif date_mode == "Last 30 Days":

            start_date_filter = (
                today - timedelta(days=29)
            )

            end_date_filter = today

        elif date_mode == "This Month":

            start_date_filter = date(
                today.year,
                today.month,
                1
            )

            end_date_filter = today

        elif date_mode == "Custom":

            col1, col2 = st.columns(2)

            with col1:

                start_date_filter = st.date_input(
                    "Start Date",
                    value=(
                        today
                        - timedelta(days=30)
                    )
                )

            with col2:

                end_date_filter = st.date_input(
                    "End Date",
                    value=today
                )

        # --------------------------------------------------
        # OTHER FILTERS
        # --------------------------------------------------

        col1, col2, col3 = st.columns(3)

        with col1:

            researcher_filter = st.selectbox(
                "Researcher",
                ["All"] + researcher_options
            )

        with col2:

            status_filter = st.selectbox(
                "Status",
                [
                    "All",
                    "Successful",
                    "Partial Success",
                    "Failed",
                    "In Progress",
                    "Not specified"
                ]
            )

        with col3:

            solvent_filter = st.selectbox(
                "Solvent",
                ["All"] + solvent_options
            )

        # --------------------------------------------------
        # APPLY FILTERS
        # --------------------------------------------------

        filtered_search_experiments = (
            filter_experiments_for_search(
                experiments_for_search,
                start_date_filter,
                end_date_filter,
                researcher_filter,
                status_filter,
                solvent_filter
            )
        )

        st.info(
            f"{len(filtered_search_experiments)} "
            f"experiment(s) match the current filters."
        )

        filter_summary = []

        if date_mode != "All Time":
            filter_summary.append(
                f"Date: {date_mode}"
            )

        if researcher_filter != "All":
            filter_summary.append(
                f"Researcher: {researcher_filter}"
            )

        if status_filter != "All":
            filter_summary.append(
                f"Status: {status_filter}"
            )

        if solvent_filter != "All":
            filter_summary.append(
                f"Solvent: {solvent_filter}"
            )

        if filter_summary:

            st.caption(
                "Current scope: "
                + " | ".join(filter_summary)
            )

        st.divider()

        st.markdown("### Ask your laboratory")

        research_question = st.text_area(
            "Question",
            placeholder=(
                "e.g. Which experiment had the highest yield? "
                "Compare my catalyst loading experiments."
            ),
            height=100
        )

        st.caption(
            "Example questions: "
            "Find all experiments using THF • "
            "Summarize failed experiments • "
            "Compare catalyst loading results • "
            "What should I try next?"
        )

        if st.button(
            "🧠 Ask LabFlow",
            type="primary",
            use_container_width=True
        ):

            if not research_question.strip():

                st.warning(
                    "Please enter a research question."
                )

            elif not filtered_search_experiments:

                st.warning(
                    "No experiments match the current filters."
                )

            else:

                with st.spinner(
                    "Searching experimental memory..."
                ):

                    try:

                        research_answer = (
                            answer_research_question(
                                research_question,
                                filtered_search_experiments
                            )
                        )

                        st.session_state[
                            "research_memory_answer"
                        ] = research_answer

                        st.session_state[
                            "research_memory_question"
                        ] = research_question

                    except Exception as e:

                        st.error(
                            f"Research search failed: {e}"
                        )

        if (
            "research_memory_answer"
            in st.session_state
        ):

            st.divider()

            st.markdown("### LabFlow Answer")

            st.write(
                st.session_state[
                    "research_memory_answer"
                ]
            )

            answer_text = st.session_state[
                "research_memory_answer"
            ]

            referenced_ids = extract_experiment_ids(
                answer_text
            )

            if referenced_ids:

                st.divider()

                st.markdown(
                    "### 🔗 Referenced Experiments"
                )

                st.caption(
                    "Experiments mentioned in the AI response."
                )

                for experiment_id in referenced_ids:

                    referenced_experiment = (
                        get_experiment_by_id(
                            experiment_id
                        )
                    )

                    if referenced_experiment is None:
                        continue

                    (
                        ref_id,
                        ref_name,
                        ref_researcher,
                        ref_date,
                        ref_starting_material,
                        ref_sm_mw,
                        ref_sm_mass,
                        ref_sm_mmol,
                        ref_reagents_json,
                        ref_solvent,
                        ref_solvent_volume,
                        ref_temperature,
                        ref_reaction_time,
                        ref_yield,
                        ref_observation,
                        ref_objective,
                        ref_status,
                        ref_key_result,
                        ref_next_step,
                        ref_cost,
                        ref_ai_procedure
                    ) = referenced_experiment

                    with st.expander(
                        f"EXP-{ref_id:04d} — {ref_name}"
                    ):

                        col1, col2, col3 = st.columns(3)

                        with col1:
                            st.write(
                                f"**Date:** {ref_date}"
                            )

                        with col2:
                            st.write(
                                f"**Yield:** "
                                f"{ref_yield:.1f}%"
                                if ref_yield is not None
                                else "**Yield:** Not recorded"
                            )

                        with col3:
                            st.write(
                                f"**Status:** "
                                f"{ref_status or 'Not specified'}"
                            )

                        st.write(
                            f"**Starting Material:** "
                            f"{ref_starting_material}"
                        )

                        st.write(
                            f"**Solvent:** "
                            f"{ref_solvent or 'Not recorded'}"
                        )

                        if ref_objective:
                            st.write(
                                f"**Objective:** "
                                f"{ref_objective}"
                            )

                        if ref_key_result:
                            st.write(
                                f"**Key Result:** "
                                f"{ref_key_result}"
                            )

                        if ref_next_step:
                            st.write(
                                f"**Next Step:** "
                                f"{ref_next_step}"
                            )

if page == "📦 Inventory":

    st.header("📦 Chemical Inventory")

    st.caption(
        "Track chemical stock, purchasing thresholds and laboratory consumption."
    )

    st.divider()

    # --------------------------------------------------
    # ADD NEW INVENTORY ITEM
    # --------------------------------------------------

    st.subheader("Add Inventory Item")

    col1, col2, col3 = st.columns(3)

    with col1:
        inventory_name = st.text_input(
            "Chemical Name",
            placeholder="e.g. THF"
        )

    with col2:
        inventory_category = st.selectbox(
            "Category",
            [
                "Solvent",
                "Reagent",
                "Catalyst",
                "Base",
                "Ligand",
                "Additive",
                "Other"
            ]
        )

    with col3:
        inventory_unit = st.selectbox(
            "Stock Unit",
            [
                "mL",
                "g"
            ]
        )

    col1, col2, col3 = st.columns(3)

    with col1:
        current_stock = st.number_input(
            "Current Stock",
            min_value=0.0,
            value=0.0,
            step=1.0
        )

    with col2:
        minimum_stock = st.number_input(
            "Low Stock Threshold",
            min_value=0.0,
            value=0.0,
            step=1.0
        )

    with col3:
        inventory_cost = st.number_input(
            "Cost per Unit (HKD)",
            min_value=0.0,
            value=0.0,
            step=0.1
        )

    col1, col2 = st.columns(2)

    with col1:
        supplier = st.text_input(
            "Supplier",
            placeholder="e.g. Sigma-Aldrich"
        )

    with col2:
        storage_location = st.text_input(
            "Storage Location",
            placeholder="e.g. Solvent Cabinet A"
        )

    if st.button(
        "➕ Add to Inventory",
        type="primary",
        use_container_width=True
    ):

        if not inventory_name:

            st.error(
                "Please enter a chemical name."
            )

        else:

            add_inventory_item(
                inventory_name,
                inventory_category,
                inventory_unit,
                current_stock,
                minimum_stock,
                inventory_cost,
                supplier,
                storage_location
            )

            st.success(
                f"{inventory_name} added to inventory."
            )

            st.rerun()

    st.divider()

    # --------------------------------------------------
    # CURRENT INVENTORY
    # --------------------------------------------------

    st.subheader("Current Inventory")

    inventory = get_inventory()

    if not inventory:

        st.info(
            "No inventory items have been added yet."
        )

    else:

        inventory_data = []

        for item in inventory:

            if item[4] <= item[5]:

                stock_status = "⚠️ Low Stock"

            else:

                stock_status = "✅ In Stock"

            inventory_data.append({
                "ID": item[0],
                "Chemical": item[1],
                "Category": item[2],
                "Unit": item[3],
                "Current Stock": item[4],
                "Minimum Stock": item[5],
                "Status": stock_status,
                "Cost / Unit (HKD)": item[6],
                "Supplier": item[7],
                "Location": item[8],
                "Last Updated": item[9]
            })

        st.dataframe(
            inventory_data,
            use_container_width=True,
            hide_index=True
        )

# --------------------------------------------------
# LAB MANAGER PAGE
# --------------------------------------------------

if page == "⚙️ Lab Manager":

    st.header("⚙️ Lab Manager Dashboard")

    st.caption(
        "Monitor laboratory activity, costs and resource usage."
    )

    st.divider()

    response = (
        supabase
        .table("experiments")
        .select(
            "id, "
            "experiment_name, "
            "researcher, "
            "experiment_date, "
            "solvent, "
            "solvent_volume, "
            "estimated_total_cost"
        )
        .order(
            "experiment_date",
            desc=True
        )
        .order(
            "id",
            desc=True
        )
        .execute()
    )

    manager_experiments = [
        (
            exp["id"],
            exp["experiment_name"],
            exp["researcher"],
            exp["experiment_date"],
            exp["solvent"],
            exp["solvent_volume"],
            exp["estimated_total_cost"]
        )
        for exp in response.data
    ]

    if not manager_experiments:

        st.info("No experiment data available yet.")

    else:

        today = date.today()

        # --------------------------------------------------
        # FILTER THIS MONTH
        # --------------------------------------------------

        monthly_experiments = []

        for exp in manager_experiments:

            try:
                exp_date = datetime.strptime(
                    exp[3],
                    "%Y-%m-%d"
                ).date()

            except:
                continue

            if (
                exp_date.year == today.year
                and exp_date.month == today.month
            ):
                monthly_experiments.append(exp)

        # --------------------------------------------------
        # SUMMARY METRICS
        # --------------------------------------------------

        total_experiments = len(monthly_experiments)

        costs = [
            exp[6] or 0
            for exp in monthly_experiments
        ]

        total_cost = sum(costs)

        average_cost = (
            total_cost / total_experiments
            if total_experiments > 0
            else 0
        )

        solvent_usage = {}

        for exp in monthly_experiments:

            solvent = exp[4]
            volume = exp[5] or 0

            if solvent:

                solvent_usage[solvent] = (
                    solvent_usage.get(solvent, 0)
                    + volume
                )

        if solvent_usage:

            most_used_solvent = max(
                solvent_usage,
                key=solvent_usage.get
            )

        else:

            most_used_solvent = "N/A"

        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.metric(
                "Experiments This Month",
                total_experiments
            )

        with col2:
            st.metric(
                "Estimated Monthly Cost",
                f"HK${total_cost:.2f}"
            )

        with col3:
            st.metric(
                "Average Cost / Experiment",
                f"HK${average_cost:.2f}"
            )

        with col4:
            st.metric(
                "Most Used Solvent",
                most_used_solvent
            )

        st.divider()

        # --------------------------------------------------
        # SOLVENT USAGE
        # --------------------------------------------------

        st.subheader("Solvent Usage")

        if solvent_usage:

            solvent_chart_data = {
                "Solvent": list(solvent_usage.keys()),
                "Volume (mL)": list(solvent_usage.values())
            }

            st.bar_chart(
                solvent_chart_data,
                x="Solvent",
                y="Volume (mL)"
            )

        else:

            st.info(
                "No solvent usage data available for this month."
            )

        st.divider()

        # --------------------------------------------------
        # EXPERIMENT COST BREAKDOWN
        # --------------------------------------------------

        st.subheader("Experiment Cost Breakdown")

        cost_data = []

        for exp in monthly_experiments:

            cost_data.append({
                "Experiment ID": f"EXP-{exp[0]:04d}",
                "Experiment": exp[1],
                "Researcher": exp[2],
                "Date": exp[3],
                "Estimated Cost (HKD)": exp[6] or 0
            })

        st.dataframe(
            cost_data,
            use_container_width=True,
            hide_index=True
        )

        st.divider()

        # --------------------------------------------------
        # COST CHART
        # --------------------------------------------------

        st.subheader("Cost by Experiment")

        if cost_data:

            st.bar_chart(
                cost_data,
                x="Experiment",
                y="Estimated Cost (HKD)"
            )